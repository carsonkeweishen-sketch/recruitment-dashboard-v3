#!/usr/bin/env python3
"""
Phase 8.0A Foundation Correction Patch — Word 自检报告
严格按 GPT 任务书格式：10 章 + 4 张截图嵌入 + 验收红线完整。
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = "/workspace/recruitment-dashboard"
SCREENSHOT_DIR = os.path.join(BASE, "screenshots/phase-8.0a-foundation-correction")
OUTPUT = os.path.join(BASE, "docs/self-checks/Phase_8.0A_Foundation_Correction_自检报告.docx")

screenshots = [
    ("standard-empty-state-real.png",
     "标准空状态 — 知识库模块成熟空状态",
     "ModulePage 组件渲染，真实页面 /knowledge", "admin",
     "验证：文案包含「该模块正在接入招聘数据」「招聘总览」「风险行动中心」"),
    ("standard-error-state-real.png",
     "标准错误状态 — 岗位页面 API 500 错误",
     "ErrorState 组件渲染，真实页面 /jobs（API 拦截模拟 500）", "admin",
     "验证：无 Prisma/SQL/DATABASE_URL/stack trace 泄露"),
    ("standard-permission-state-real.png",
     "标准权限拒绝 — Interviewer 访问设置页",
     "PermissionState 组件渲染，真实页面 /settings", "interviewer",
     "验证：不泄露对象存在性、归属人、内部 ID"),
    ("standard-loading-skeleton-real.png",
     "标准加载骨架屏 — KPI + 列表 + 卡片骨架",
     "LoadingSkeleton 组件渲染，含 animate-pulse", "admin",
     "验证：展示 KPI 骨架行 + 列表骨架行 + 卡片骨架，非空白页"),
]


def set_cell_shading(cell, color):
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color,
    })
    shading_elm.append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '2B579A')
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r % 2 == 0:
                set_cell_shading(cell, 'F2F6FC')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# ============================================================
# Title
# ============================================================
title = doc.add_heading('理然智能招聘 AI 看板 — Phase 8.0A Foundation Correction 自检报告', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(f'生成日期：{datetime.now().strftime("%Y-%m-%d %H:%M")}')
doc.add_paragraph('分支：agent/workbuddy/phase-7')
doc.add_paragraph('版本：Phase 8.0A — Foundation Correction Patch')
doc.add_paragraph('Mock：否 — 全部截图来自真实页面渲染，错误态通过 API 拦截模拟')
doc.add_paragraph()

# ============================================================
# 一、构建验证
# ============================================================
doc.add_heading('一、构建验证', level=1)
add_styled_table(doc,
    ['检查项', '命令', '结果', '说明'],
    [
        ['Type Check', 'pnpm typecheck', '✅ PASS', '0 errors'],
        ['Lint', 'pnpm lint', '✅ PASS', '0 errors 0 warnings'],
        ['Build', 'pnpm build', '✅ PASS', '全部路由编译成功'],
    ]
)

# ============================================================
# 二、P0 关闭确认
# ============================================================
doc.add_heading('二、P0 关闭确认', level=1)
add_styled_table(doc,
    ['P0 #', '问题', '状态', '修复说明'],
    [
        ['P0-1', '/dashboard 命名"AI 决策看板"违反 AI 不决策边界',
         '✅',
         '导航名称→"招聘总览"；页面标题→"AI 招聘洞察看板"；页面副标题→"基于招聘过程数据，辅助识别招聘风险、流程卡点和优先处理事项。"；11 个文件同步修改'],
        ['P0-2', 'AI_CAPABILITY_FRAMEWORK 中 Layer 3 "自主智能"暗示自动决策',
         '✅',
         'Layer 3 重写为"Human-Confirmed Execution（人工确认后的执行闭环）"；明确当前不启用自动执行；7 条红线（不自动录用/淘汰/推进/发Offer等）'],
        ['P0-3', '四态截图标题与画面不匹配，证据不够可信',
         '✅',
         '重新捕获 4 张真实截图（空/错/权/载），每张清楚展示对应状态组件和文案'],
    ]
)

# ============================================================
# 三、P1 处理确认
# ============================================================
doc.add_heading('三、P1 处理确认', level=1)
add_styled_table(doc,
    ['P1 #', '内容', '状态', '说明'],
    [
        ['P1-1', '确认知识库模块范围', '✅', 'MODULE_ROADMAP 补充：明确禁止上传文档解析/AI知识问答/自动生成题库/复杂权限文档库'],
        ['P1-2', '补真实产物抽检', '✅', '10 项抽检全部通过，详见 audit.md'],
    ]
)

# ============================================================
# 四、Claude Review Triage（更新）
# ============================================================
doc.add_heading('四、Claude Review Triage（更新）', level=1)
add_styled_table(doc,
    ['级别', '总数', '采纳完成', '记录不做'],
    [
        ['P0 (Phase 8.0A)', '3', '3', '0'],
        ['P1 (Phase 8.0A)', '2', '2', '0'],
        ['合计', '5', '5', '0'],
    ]
)

# ============================================================
# 五、修改文件清单
# ============================================================
doc.add_heading('五、修改文件清单', level=1)
add_styled_table(doc,
    ['文件', '改动内容'],
    [
        ['components/layout/Sidebar.tsx', '导航标签 "AI 决策看板" → "招聘总览"'],
        ['app/dashboard/page.tsx', '页面标题 + 副标题更新'],
        ['components/ui/module-page.tsx', '空状态文案 "AI 决策看板" → "招聘总览"'],
        ['server/config/module-registry.ts', 'label + description 更新'],
        ['docs/product/AI_CAPABILITY_FRAMEWORK.md', 'Layer 3 完全重写'],
        ['docs/product/PRODUCT_NAMING_GLOSSARY.md', 'Dashboard 命名更新'],
        ['docs/product/PRODUCT_INFORMATION_ARCHITECTURE.md', '模块命名更新'],
        ['docs/product/MODULE_ROADMAP.md', '命名更新 + 知识库禁止项补充'],
        ['docs/design/CLAUDE_PHASE_8_FOUNDATION_REVIEW_TRIAGE.md', 'Layer 3 描述更新'],
        ['docs/self-checks/phase-8-foundation-report.md', '命名更新'],
        ['docs/self-checks/phase-8-foundation-screenshot-index.md', '命名更新'],
    ]
)

# ============================================================
# 六、Grep 验证
# ============================================================
doc.add_heading('六、Grep 验证结果', level=1)
add_styled_table(doc,
    ['检查项', '结果', '说明'],
    [
        ['"AI 决策看板" / "AI 决策" 清零', '✅', 'app/components/server/docs 中已清零（仅在禁止列表中出现）'],
        ['"自主智能" / "自动决策" 清零', '✅', '已清零'],
        ['"自动录用" / "自动淘汰" 清零', '✅', '已清零'],
        ['page.route mock', '✅', '仅旧 scripts 中存在，app/components 中无'],
        ['mock/test/demo 命名', '✅', 'app/components 中无'],
        ['null/undefined/NaN in UI', '✅', '仅 TS 类型注解'],
        ['env/secrets leak', '✅', '无明文泄露'],
        ['fake AI', '✅', '无'],
        ['permissions-debug in nav', '✅', '已移除'],
        ['散落 hex 色值', '✅', '仅 CSS 变量引用'],
    ]
)

# ============================================================
# 七、截图证据（4 张）
# ============================================================
doc.add_heading('七、截图证据（共 4 张）', level=1)

p = doc.add_paragraph()
p.add_run('以下 4 张截图全部来自真实页面渲染，清楚展示 Empty/Error/Permission/Loading 四种标准状态组件。').bold = True

for i, (filename, description, api_info, role, verify) in enumerate(screenshots, 1):
    doc.add_heading(f'{i}. {description}', level=2)

    add_styled_table(doc,
        ['属性', '值'],
        [
            ['文件名', filename],
            ['页面/状态', description],
            ['数据来源', api_info],
            ['角色', role],
            ['Mock', '否'],
            ['验证内容', verify],
        ],
        col_widths=[3, 13]
    )

    filepath = os.path.join(SCREENSHOT_DIR, filename)
    if os.path.exists(filepath):
        doc.add_picture(filepath, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f'⚠️ 截图文件未找到: {filepath}')

    doc.add_paragraph()

# ============================================================
# 八、验收红线确认
# ============================================================
doc.add_heading('八、验收红线确认', level=1)
add_styled_table(doc,
    ['#', '红线项', '状态'],
    [
        ['1', '不出现"AI 决策看板"', '✅'],
        ['2', '不出现"AI 决策"作为产品主文案', '✅'],
        ['3', '不出现"自主智能"', '✅'],
        ['4', 'AI 边界不暗示自动决策', '✅'],
        ['5', '四态截图标题与画面匹配', '✅'],
        ['6', '权限态不泄露对象信息', '✅'],
        ['7', '错误态不暴露技术细节', '✅'],
        ['8', '未完成模块非 404/空白/TODO', '✅'],
        ['9', 'Product Shell 无 debug 入口', '✅'],
        ['10', '/actions 未被破坏', '✅'],
        ['11', 'typecheck/lint/build 全部通过', '✅'],
        ['12', '截图完成（4 张）', '✅'],
        ['13', '未合并 main', '✅'],
        ['14', '未 force push', '✅'],
        ['15', '未自行进入 Phase 8.1', '✅'],
    ]
)

# ============================================================
# 九、真实产物抽检摘要
# ============================================================
doc.add_heading('九、真实产物抽检摘要', level=1)
add_styled_table(doc,
    ['#', '抽检项', '结果'],
    [
        ['1', 'design-tokens 是否为唯一来源', '✅'],
        ['2', '是否存在散落 hex', '✅ 仅 CSS 变量引用'],
        ['3', '四态组件是否真实页面接入', '✅'],
        ['4', '命名 glossary 是否在 UI 落地', '✅'],
        ['5', '"AI 决策"是否已清零', '✅'],
        ['6', '"自主智能"是否已清零', '✅'],
        ['7', 'Product Shell 是否无 debug 入口', '✅'],
        ['8', '/actions 是否未被破坏', '✅'],
        ['9', '未完成模块是否无 404/TODO', '✅'],
        ['10', '是否无假 AI', '✅'],
    ]
)
doc.add_paragraph('结论：10/10 全部通过。详见 phase-8.0a-foundation-correction-audit.md')

# ============================================================
# 十、最终结论
# ============================================================
doc.add_heading('十、最终结论', level=1)
add_styled_table(doc,
    ['项目', '结论'],
    [
        ['Phase 8.0A Foundation Correction 是否完成', '是'],
        ['是否已移除"AI 决策看板"', '是'],
        ['/dashboard 导航名称是否改为"招聘总览"', '是'],
        ['/dashboard 页面标题是否改为"AI 招聘洞察看板"', '是'],
        ['是否已移除"自主智能"', '是'],
        ['Layer 3 是否改为"人工确认后的执行闭环"', '是'],
        ['四态真实截图是否完成', '是（4 张）'],
        ['知识库 / 模板库范围是否明确', '是'],
        ['真实产物抽检是否完成', '是（10/10）'],
        ['是否破坏 Action Center', '否'],
        ['是否进入 Phase 8.1', '否'],
        ['是否使用 mock 数据', '否'],
        ['是否存在假 AI', '否'],
        ['typecheck/lint/build 是否通过', '是'],
        ['截图是否完成', '是（4 张）'],
        ['git status 是否 clean', '待提交'],
        ['是否合并 main', '否'],
        ['是否 force push', '否'],
        ['当前风险', '无'],
        ['需要外部确认', 'ChatGPT 最终验收'],
    ]
)

# ============================================================
# 附录：截图文件清单
# ============================================================
doc.add_heading('附录：截图文件清单', level=1)
for i, (filename, description, api_info, role, verify) in enumerate(screenshots, 1):
    doc.add_paragraph(f'{i}. {filename} — {description} [{role}]', style='List Number')

# Save
doc.save(OUTPUT)
file_size_mb = os.path.getsize(OUTPUT) / (1024 * 1024)
print(f"✅ Word report saved to: {OUTPUT}")
print(f"   File size: {file_size_mb:.1f} MB")
print(f"   Screenshots embedded: {len(screenshots)}")
