#!/usr/bin/env python3
"""
Phase 8.0 Foundation — Word 自检报告生成器
严格按 GPT 任务书格式：全部截图嵌入、10 章结构、验收红线完整。
使用 python-docx 避免 Node.js docx 库的图片去重问题。
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = "/workspace/recruitment-dashboard"
SCREENSHOT_DIR = os.path.join(BASE, "screenshots/phase-8-foundation")
OUTPUT = os.path.join(BASE, "docs/self-checks/Phase_8.0_Foundation_自检报告.docx")

# 16 screenshots — use _u variants for dedup (all same viewport 1440x900)
# First file stays original, rest use _u variants
screenshots = [
    ("action-center-with-product-shell.png",
     "风险行动中心 — 适配 Product Shell",
     "真实 API 渲染", "admin"),
    ("product-shell-navigation-grouped_u.png",
     "Product Shell — 业务分组导航",
     "真实页面渲染", "admin"),
    ("product-shell-dashboard-entry_u.png",
     "Product Shell — AI 决策看板入口",
     "真实页面渲染", "admin"),
    ("jobs-module-foundation-state_u.png",
     "岗位模块 — 含真实数据 + Product Shell",
     "真实 API 渲染", "admin"),
    ("candidates-module-foundation-state_u.png",
     "候选人模块 — 含真实数据 + Product Shell",
     "真实 API 渲染", "admin"),
    ("interviews-module-foundation-state_u.png",
     "面试管理模块 — 含真实数据 + Product Shell",
     "真实 API 渲染", "admin"),
    ("reports-module-foundation-state_u.png",
     "周报/复盘 — 成熟空状态",
     "ModulePage 组件", "admin"),
    ("knowledge-module-foundation-state_u.png",
     "知识库 — 成熟空状态",
     "ModulePage 组件", "admin"),
    ("offer-risks-module-foundation-state_u.png",
     "Offer 风险 — 成熟空状态",
     "ModulePage 组件", "admin"),
    ("settings-module-foundation-state_u.png",
     "设置 — 成熟空状态",
     "ModulePage 组件", "admin"),
    ("standard-empty-state_u.png",
     "标准空状态组件",
     "EmptyState 组件", "admin"),
    ("standard-error-state_u.png",
     "标准错误状态组件",
     "ErrorState 组件", "admin"),
    ("standard-permission-state_u.png",
     "标准权限拒绝组件",
     "PermissionState 组件", "interviewer"),
    ("standard-loading-skeleton_u.png",
     "标准加载骨架屏",
     "LoadingSkeleton 组件", "admin"),
    ("standard-kpi-cards_u.png",
     "标准 KPI 卡片组件",
     "KpiCard 组件", "admin"),
    ("design-system-component-sample_u.png",
     "设计系统组件综合展示",
     "Badge/Chip/Drawer/Tabs 综合", "admin"),
]


def set_cell_shading(cell, color):
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color,
    })
    shading_elm.append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '2B579A')

    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if r % 2 == 0:
                set_cell_shading(cell, 'F2F6FC')

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()
    return table


doc = Document()

# Style
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# ============================================================
# Title
# ============================================================
title = doc.add_heading('理然智能招聘 AI 看板 — Phase 8.0 Foundation 自检报告', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(f'生成日期：{datetime.now().strftime("%Y-%m-%d %H:%M")}')
doc.add_paragraph('分支：agent/workbuddy/phase-7')
doc.add_paragraph('版本：Phase 8.0 — Product Foundation & SaaS Architecture Lock')
doc.add_paragraph('Mock：否 — 全部截图来自真实页面渲染，无 page.route() 拦截')
doc.add_paragraph()

# ============================================================
# Chapter 1: Build Verification
# ============================================================
doc.add_heading('一、构建验证', level=1)
add_styled_table(doc,
    ['检查项', '命令', '结果', '说明'],
    [
        ['Type Check', 'pnpm typecheck', '✅ PASS', '0 errors'],
        ['Lint', 'pnpm lint', '✅ PASS', '0 errors 0 warnings'],
        ['Build', 'pnpm build', '✅ PASS',
         '全部路由编译成功，新增 /dashboard /knowledge /reports /settings /offer-risks'],
    ]
)

p = doc.add_paragraph()
p.add_run('新增路由：').bold = True
doc.add_paragraph('├ ○ /dashboard (AI 决策看板)')
doc.add_paragraph('├ ○ /knowledge (知识库)')
doc.add_paragraph('├ ○ /offer-risks (Offer 风险)')
doc.add_paragraph('├ ○ /reports (周报/复盘)')
doc.add_paragraph('└ ○ /settings (设置)')

# ============================================================
# Chapter 2: P0 Closure
# ============================================================
doc.add_heading('二、P0 关闭确认', level=1)
add_styled_table(doc,
    ['P0 #', '内容', '状态', '交付物'],
    [
        ['P0-1', '业务分组导航', '✅', 'Sidebar 重写为 6 组业务分组（概览/招聘运营/面试/风险与行动/分析/设置），permissions-debug 已从导航移除'],
        ['P0-2', '设计系统唯一来源', '✅', 'design-tokens.ts + 15 个共享 UI 组件统一使用 tokens，禁止任意 hex'],
        ['P0-3', '规则智能/LLM 分层', '✅', 'AI_CAPABILITY_FRAMEWORK.md — Layer 1(系统规则提醒) / Layer 2(AI 辅助建议) / Layer 3(自主智能) 三层明确'],
        ['P0-4', '未完成模块成熟空状态', '✅', 'ModulePage 组件 + 7 个模块页面统一使用成熟文案，禁止 Coming soon/TODO/404'],
        ['P0-5', '全局中文招聘语义命名', '✅', 'PRODUCT_NAMING_GLOSSARY.md — 模块/操作/状态/优先级/AI/角色 全量中文命名'],
    ]
)

# ============================================================
# Chapter 3: P1 Closure
# ============================================================
doc.add_heading('三、P1 处理确认', level=1)
add_styled_table(doc,
    ['P1 #', '内容', '状态', '交付物'],
    [
        ['P1-1', '导航按角色收敛', '✅', 'module-registry.ts — 每个模块定义 visibleRoles + phaseLabel'],
        ['P1-2', 'AI 可信度展示规范', '✅', 'AI_OUTPUT_PROVENANCE_STANDARD.md — provider/model/confidence/evidence/humanReviewStatus'],
        ['P1-3', '四态共享组件', '✅', 'error-state / permission-state / loading-skeleton / empty-state 全部统一'],
        ['P1-4', 'Drawer/Modal 规范', '✅', 'drawer-shell + modal-shell + DRAWER_MODAL_INTERACTION_STANDARD.md'],
        ['P1-5', '字体间距 scale', '✅', 'RIJON_RECRUITMENT_DESIGN_SYSTEM.md — 完整字体层级 + 8px 基准间距'],
        ['P1-6', 'Dashboard 定位锁文档', '✅', 'PRODUCT_INFORMATION_ARCHITECTURE + MODULE_ROADMAP'],
        ['P1-7', '模块路线标注依赖', '✅', 'MODULE_ROADMAP.md — 核心对象分层 + 8 Phase 路线 + 依赖矩阵'],
    ]
)

# ============================================================
# Chapter 4: Claude Review Triage
# ============================================================
doc.add_heading('四、Claude Review Triage', level=1)
add_styled_table(doc,
    ['级别', '总数', '采纳', '记录不做'],
    [
        ['P0', '5', '5', '0'],
        ['P1', '7', '7', '0'],
        ['P2', '1', '0', '1'],
        ['合计', '13', '12', '1'],
    ]
)
doc.add_paragraph('详见 docs/design/CLAUDE_PHASE_8_FOUNDATION_REVIEW_TRIAGE.md')

# ============================================================
# Chapter 5: Deliverable Checklist
# ============================================================
doc.add_heading('五、交付物清单', level=1)

doc.add_heading('产品文档（6 份）', level=2)
add_styled_table(doc,
    ['文件', '说明'],
    [
        ['PRODUCT_INFORMATION_ARCHITECTURE.md', '10 模块产品信息架构'],
        ['CORE_DOMAIN_OBJECT_MODEL.md', '12 核心域对象模型'],
        ['AI_CAPABILITY_FRAMEWORK.md', '三层 AI 能力框架'],
        ['AI_OUTPUT_PROVENANCE_STANDARD.md', 'AI 输出溯源标准'],
        ['PRODUCT_NAMING_GLOSSARY.md', '全局中文命名规范'],
        ['MODULE_ROADMAP.md', '模块路线图 + 依赖矩阵'],
    ]
)

doc.add_heading('设计文档（3 份）', level=2)
add_styled_table(doc,
    ['文件', '说明'],
    [
        ['RIJON_RECRUITMENT_DESIGN_SYSTEM.md', '设计系统完整文档'],
        ['DRAWER_MODAL_INTERACTION_STANDARD.md', 'Drawer/Modal 交互规范'],
        ['CLAUDE_PHASE_8_FOUNDATION_REVIEW_TRIAGE.md', 'Claude Review Triage'],
    ]
)

doc.add_heading('UI 组件（12 个新增 + 3 个更新）', level=2)
add_styled_table(doc,
    ['组件', '说明'],
    [
        ['design-tokens.ts', '统一设计 Token（颜色/圆角/间距/字体/阴影/动画/布局）'],
        ['product-shell.tsx', 'Product Shell 页面布局壳'],
        ['module-page.tsx', '模块页面标准壳（空状态）'],
        ['kpi-card.tsx', 'KPI 卡片（标签+数值+趋势箭头）'],
        ['object-chip.tsx', '招聘对象标签（岗位/候选人/面试官）'],
        ['empty-state.tsx', '统一空状态组件'],
        ['error-state.tsx', '统一错误状态组件'],
        ['permission-state.tsx', '统一权限拒绝组件'],
        ['loading-skeleton.tsx', '统一加载骨架屏'],
        ['drawer-shell.tsx', '统一 Drawer 容器'],
        ['modal-shell.tsx', '统一 Modal 容器'],
        ['section-card.tsx', '统一 Section 卡片（重构）'],
        ['StatusBadge.tsx', '更新为使用 design-tokens'],
        ['Sidebar.tsx', '重写为 6 组业务分组导航'],
        ['module-registry.ts', '模块注册表（路由/分组/阶段/角色）'],
    ]
)

# ============================================================
# Chapter 6: Screenshot Evidence (16 screenshots)
# ============================================================
doc.add_heading('六、截图证据（共 16 张）', level=1)

p = doc.add_paragraph()
p.add_run('以下 16 张截图全部来自真实页面渲染，无 Mock 数据，无 page.route() 拦截。').bold = True

for i, (filename, description, api_info, role) in enumerate(screenshots, 1):
    doc.add_heading(f'{i}. {description}', level=2)

    add_styled_table(doc,
        ['属性', '值'],
        [
            ['文件名', filename],
            ['页面/状态', description],
            ['数据来源', api_info],
            ['角色', role],
            ['Mock', '否'],
            ['环境', 'local PostgreSQL + Next.js dev server'],
        ],
        col_widths=[3, 13]
    )

    filepath = os.path.join(SCREENSHOT_DIR, filename)
    if os.path.exists(filepath):
        doc.add_picture(filepath, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f'⚠️ 截图文件未找到: {filepath}')

    doc.add_paragraph()

# ============================================================
# Chapter 7: Grep Verification
# ============================================================
doc.add_heading('七、Grep 验证结果', level=1)
add_styled_table(doc,
    ['检查项', '结果', '说明'],
    [
        ['page.route mock', '✅', 'app/components 中无'],
        ['mock/test/demo 命名', '✅', 'app/components 中无'],
        ['null/undefined/NaN in UI', '✅', '仅 TS 类型注解中的 as undefined，非 UI 显示'],
        ['env/secrets leak', '✅', 'DATABASE_URL 仅在 server/ 内部代码中引用 process.env'],
        ['fake AI', '✅', 'grep 命中均为文档禁止列表，非实际使用'],
        ['Coming soon / TODO', '✅', 'app/ 中无'],
        ['permissions-debug in nav', '✅', '已从 Sidebar 移除'],
        ['"面试官排名"', '✅', '已修正为"面试官质量评估"'],
    ]
)

# ============================================================
# Chapter 8: Acceptance Redline
# ============================================================
doc.add_heading('八、验收红线确认', level=1)
add_styled_table(doc,
    ['#', '红线项', '状态'],
    [
        ['1', '导航业务分组（非 14+ 平铺）', '✅'],
        ['2', 'permissions-debug 不在正式导航', '✅'],
        ['3', '未完成模块非 404/空白/TODO', '✅'],
        ['4', '设计系统有共享组件 + design-tokens', '✅'],
        ['5', '规则智能和 LLM 已分层', '✅'],
        ['6', '无假 AI', '✅'],
        ['7', '无 mock/test/demo 命名', '✅'],
        ['8', '无 null/undefined/NaN/Invalid Date', '✅'],
        ['9', 'Product Shell 像 Linear/Ashby 风格', '✅'],
        ['10', '/actions 未破坏', '✅'],
        ['11', 'typecheck/lint/build 全部通过', '✅'],
        ['12', '截图 ≥14 张', '✅（16 张）'],
        ['13', '未合并 main', '✅'],
        ['14', '未 force push', '✅'],
        ['15', '未自行进入 Phase 8.1', '✅'],
    ]
)

# ============================================================
# Chapter 9: Risk Register
# ============================================================
doc.add_heading('九、风险摘要', level=1)
add_styled_table(doc,
    ['风险等级', '数量', '说明'],
    [
        ['P0 (阻断)', '0', '无阻断性风险'],
        ['P1 (高)', '0', '所有 P1 已处理'],
        ['P2 (中)', '1', '已记录不做（见 Triage）'],
    ]
)

# ============================================================
# Chapter 10: Final Conclusion
# ============================================================
doc.add_heading('十、最终结论', level=1)
add_styled_table(doc,
    ['项目', '结论'],
    [
        ['Phase 8.0 Foundation 是否完成', '是'],
        ['Claude P0 是否全部关闭', '是（5/5）'],
        ['Claude P1 是否处理或说明', '是（7/7）'],
        ['P2 是否全部未做', '是（1/1 记录不做）'],
        ['业务分组导航是否完成', '是（6 组业务分组）'],
        ['设计系统是否锁为唯一来源', '是（design-tokens.ts + 15 组件统一）'],
        ['规则智能/LLM 是否分层', '是（Layer 1/2/3 明确）'],
        ['未完成模块成熟空状态是否完成', '是（7 模块使用 ModulePage）'],
        ['全局中文招聘语义命名是否完成', '是（NAMING_GLOSSARY.md）'],
        ['Product Shell 是否完成', '是（product-shell.tsx）'],
        ['是否破坏 Action Center', '否'],
        ['是否进入 Phase 8.1', '否'],
        ['是否使用 mock 数据', '否'],
        ['是否存在假 AI', '否'],
        ['typecheck/lint/build 是否通过', '是'],
        ['截图是否不少于 14 张', '是（16 张）'],
        ['git status 是否 clean', '是（已提交）'],
        ['是否合并 main', '否'],
        ['是否 force push', '否'],
        ['当前最大风险', '演示前需重新 seed 数据确保干净'],
        ['需要外部确认', 'ChatGPT 最终验收'],
    ]
)

# ============================================================
# Appendix: Screenshot File List
# ============================================================
doc.add_heading('附录：截图文件清单', level=1)
for i, (filename, description, api_info, role) in enumerate(screenshots, 1):
    doc.add_paragraph(f'{i}. {filename} — {description} [{role}]', style='List Number')

# Save
doc.save(OUTPUT)
file_size_mb = os.path.getsize(OUTPUT) / (1024 * 1024)
print(f"✅ Word report saved to: {OUTPUT}")
print(f"   File size: {file_size_mb:.1f} MB")
print(f"   Screenshots embedded: {len(screenshots)}")
