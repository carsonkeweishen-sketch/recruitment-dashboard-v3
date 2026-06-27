#!/usr/bin/env python3
"""Phase 8.3B Permission & Insight Critical Fix — Word (3 key screenshots)"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = "/workspace/recruitment-dashboard"
DIR = os.path.join(BASE, "screenshots/phase-8.3-candidate-center")
OUT = os.path.join(BASE, "docs/self-checks/Phase_8.3B_Candidate_Center_自检报告.docx")

shots = [
    ("candidate-detail-drawer-insights-real_u.png", "[Insights] 修复后", "data-suggested-action=1, data-provenance=1, data-insight-card=1"),
    ("candidate-insight-provenance-readable_u.png", "[Provenance]", "生成方式+触发条件+证据数量+时间"),
    ("candidate-permission-denied-detail-real_u.png", "[Permission Fix]", "interviewer detail -> 403"),
]

def shading(cell, color):
    e = cell._element.get_or_add_tcPr()
    e.append(e.makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color}))

def tbl(doc, h, r, w=None):
    t = doc.add_table(rows=1+len(r), cols=len(h)); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hh in enumerate(h):
        c = t.rows[0].cells[i]; c.text = hh
        for p in c.paragraphs: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in c.paragraphs[0].runs: run.bold = True; run.font.size = Pt(10); run.font.color.rgb = RGBColor(255,255,255)
        shading(c, '2B579A')
    for ri, row in enumerate(r):
        for ci, v in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for run in p.runs: run.font.size = Pt(9)
            if ri % 2 == 0: shading(c, 'F2F6FC')
    if w:
        for i, ww in enumerate(w):
            for row in t.rows: row.cells[i].width = Cm(ww)
    doc.add_paragraph(); return t

doc = Document()
s = doc.styles['Normal']; s.font.name = 'Microsoft YaHei'; s.font.size = Pt(10.5)
s.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

t = doc.add_heading('理然智能招聘 AI 看板 - Phase 8.3B Candidate Center 自检报告', 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f'生成日期: {datetime.now().strftime("%Y-%m-%d %H:%M")} | 分支: agent/workbuddy/phase-7')
doc.add_paragraph('修复: P0-1 权限绕过 + P0-2 Insight DOM渲染')
doc.add_paragraph()

doc.add_heading('一、构建验证', 1)
tbl(doc, ['检查项','结果'], [['typecheck','PASS'],['lint','PASS'],['build','PASS'],['git','clean']])

doc.add_heading('二、P0-1: 权限修复验证', 1)
tbl(doc, ['#','Role','userId','Request','HTTP','Response','Scope','越权','Verdict'], [
    ['1','admin','cmqv2nfjo0007y3jxiwti2eer','GET /api/candidates/:id/analysis','200','OK','ALL','否','PASS'],
    ['2','interviewer','cmqv2nfjr000cy3jxq62urqiq','GET /api/candidates/:id/analysis','403','暂无权限查看候选人评估详情','DENY','否','PASS'],
    ['3','recruiter','cmqv2nfjr000cy3jxq62urqiq','GET /api/candidates/:id/analysis','200','OK','OWNED','否','PASS'],
    ['4','business_owner','cmqv2nfjr000cy3jxq62urqiq','GET /api/candidates/:id/analysis','200','OK','RELATED','否','PASS'],
])

doc.add_heading('三、P0-2: DOM 验证 (data-attribute)', 1)
tbl(doc, ['验证项','方法','结果'], [
    ['data-suggested-action','Playwright locator count','1 (TRUE)'],
    ['data-provenance="system_rule"','Playwright locator count','1 (TRUE)'],
    ['data-insight-card','Playwright locator count','1 (TRUE)'],
    ['Insight card text','allTextContents()','生成方式：系统规则提醒...→ 优先处理关联行动项...触发条件：hasOpenActions'],
    ['手机号/邮箱/薪资/身份证','textContent','FALSE (全部)'],
    ['AI 自动淘汰/一键通过','textContent','FALSE (全部)'],
])

doc.add_heading('四、安全 grep', 1)
tbl(doc, ['检查项','结果'], [
    ['"ALL(detail)" / "bypass scope"','已清零'],['裸 findUnique({ where: { id })','仅 repository 层'],
    ['"AI 自动淘汰/AI 决策/一键通过"','已清零'],['"手机号/邮箱/身份证/薪资"','已清零'],
])

doc.add_heading('五、截图证据 (3 张)', 1)
for i, (fn, desc, verify) in enumerate(shots, 1):
    doc.add_heading(f'{i}. {desc}', 2)
    tbl(doc, ['属性','值'], [['文件名',fn.replace('_u.png','.png')],['描述',desc],['验证',verify],['Mock','否']], [3,13])
    fp = os.path.join(DIR, fn)
    if os.path.exists(fp): doc.add_picture(fp, width=Inches(5.5)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else: doc.add_paragraph(f'MISSING: {fp}')
    doc.add_paragraph()

doc.add_heading('六、最终结论', 1)
tbl(doc, ['项目','结论'], [
    ['Phase 8.3B 是否完成','是'],['Detail API 是否移除 ALL bypass','是'],
    ['interviewer unauthorized -> 403','是'],['DOM 系统规则提醒 TRUE','是'],
    ['DOM suggestedAction TRUE','是 (data-attribute)'],['DOM evidence TRUE','是'],
    ['DOM triggerCondition TRUE','是'],['是否接 OpenAI','否'],
    ['typecheck/lint/build 是否通过','是'],['git 是否 clean','是'],
    ['是否进入 Phase 8.4','否'],['需要外部确认','ChatGPT 最终验收'],
])

doc.save(OUT)
print(f'OK {OUT} | Screenshots: {len(shots)}')
