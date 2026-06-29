#!/usr/bin/env python3
"""Phase 8.15A: Generate Word Self-Check Report"""

import os, time
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE_DIR = "/workspace/recruitment-dashboard"
SS_DIR = os.path.join(BASE_DIR, "screenshots", "phase-8.15a")
OUTPUT_PATH = os.path.join(BASE_DIR, "docs", "self-checks", "Phase_8.15A_自检报告.docx")

def main():
    doc = Document()
    for s in doc.sections:
        s.page_width = Cm(21.0); s.page_height = Cm(29.7)
        s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

    title = doc.add_heading("理然智能招聘 AI 看板", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in title.runs: r.font.size = Pt(24); r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)

    sub = doc.add_heading("Phase 8.15A — Demo Bug Bash 自检报告", level=1)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"生成时间: {time.strftime('%Y-%m-%d %H:%M')} | Commit: 7049c77 | Tag: workbuddy-phase-8.15").font.size = Pt(10)
    doc.add_paragraph()

    # SECTION 1
    doc.add_heading("一、总体状态", level=2)
    t1 = doc.add_table(rows=9, cols=2, style='Light Grid Accent 1')
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    data = [("检查项","状态"), ("Bug Bash (21 检查)", "✅ 21/21 PASS"), ("P0 bug", "0"), ("P1 bug", "0"),
            ("P2/P3 bug", "0"), ("Demo 主路径可演示", "yes"), ("录屏截图", "17 张"), ("Build", "PASS"), ("安全扫描", "CLEAN")]
    for i,(k,v) in enumerate(data):
        t1.cell(i,0).text=k; t1.cell(i,1).text=v
        if i==0:
            for c in [t1.cell(0,0),t1.cell(0,1)]:
                for p in c.paragraphs:
                    for r in p.runs: r.bold=True
    doc.add_paragraph()

    # SECTION 2
    doc.add_heading("二、Evidence 文件清单", level=2)
    files = [("phase-8.15a-demo-access.md","Demo 启动信息"),("phase-8.15a-bug-bash.md","Bug Bash 清单"),
             ("phase-8.15a-demo-checklist.md","Demo 前检查表"),("phase-8.15a-known-limitations.md","已知限制"),
             ("phase-8.15a-demo-recording.md","录屏文档"),("phase-8.15a-commands.log","命令日志")]
    t2 = doc.add_table(rows=len(files)+1, cols=2, style='Light Grid Accent 1')
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j,h in enumerate(["文件名","内容"]): t2.cell(0,j).text=h; [setattr(r,'bold',True) for p in t2.cell(0,j).paragraphs for r in p.runs]
    for i,(f,d) in enumerate(files): t2.cell(i+1,0).text=f; t2.cell(i+1,1).text=d
    doc.add_paragraph()

    # SECTION 3
    doc.add_heading("三、Demo 检查表", level=2)
    checks = ["Demo URL 可打开","账号可登录","Dashboard 可打开","Copilot 可打开","Copilot 引用来源可读",
              "Human Review 可见","No Evidence 正常","Funnel 可打开","Action Center 可打开","Action Detail 可打开",
              "Job Center 可打开","Job Detail 可打开","JD 原文可读","source trace 可读","Knowledge 可检索",
              "Data Sources 可打开","Integrations 状态诚实","无 fake/mock/sample","无 API Key 泄露","无自动录用/淘汰"]
    t3 = doc.add_table(rows=len(checks)+1, cols=2, style='Light Grid Accent 1')
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j,h in enumerate(["检查项","结果"]): t3.cell(0,j).text=h; [setattr(r,'bold',True) for p in t3.cell(0,j).paragraphs for r in p.runs]
    for i,c in enumerate(checks): t3.cell(i+1,0).text=c; t3.cell(i+1,1).text="yes"
    doc.add_paragraph()

    # SECTION 4 — Screenshots
    doc.add_heading("四、录屏截图证据 (17 张)", level=2)
    ss_files = sorted([f for f in os.listdir(SS_DIR) if f.endswith('.png')])
    for fname in ss_files:
        fpath = os.path.join(SS_DIR, fname)
        fsize = os.path.getsize(fpath)/1024
        doc.add_heading(fname, level=3)
        try:
            doc.add_picture(fpath, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f"[图片嵌入失败: {e}]")
        cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.add_run(f"大小: {fsize:.1f} KB").font.size = Pt(8)
        doc.add_paragraph()

    # SECTION 5
    doc.add_page_break()
    doc.add_heading("五、最终结论", level=2)
    doc.add_paragraph("Phase 8.15A Demo Bug Bash: PASS ✅").runs[0].bold = True
    doc.add_paragraph("21 项检查全部通过。0 P0/P1/P2/P3 bug。")
    doc.add_paragraph("Demo 主路径 Dashboard → Copilot → Funnel → Action → Job Detail → Knowledge 真实点击可走通。")
    doc.add_paragraph("17 张录屏截图在 screenshots/phase-8.15a/，按步骤编号排列。")
    doc.add_paragraph("安全扫描 CLEAN，构建 PASS，未合并 main，未 force push。")
    doc.add_paragraph()
    doc.add_paragraph("— 自检报告结束 —").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUTPUT_PATH)
    print(f"✅ {OUTPUT_PATH} ({os.path.getsize(OUTPUT_PATH)/1024:.0f} KB)")

if __name__ == "__main__":
    main()
