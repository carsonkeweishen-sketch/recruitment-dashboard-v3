#!/usr/bin/env python3
"""Phase 8.1 Jobs Module — Word 自检报告"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = "/workspace/recruitment-dashboard"
SCREENSHOT_DIR = os.path.join(BASE, "screenshots/phase-8.1-jobs")
OUTPUT = os.path.join(BASE, "docs/self-checks/Phase_8.1_Jobs_自检报告.docx")

screenshots = [
    ("jobs-list-with-kpi_u.png", "岗位列表视图 — KPI卡片+筛选+列表", "真实 API 渲染", "admin"),
    ("jobs-detail-overview_u.png", "岗位详情 — Overview（编制进度/风险/行动项）", "GET /api/jobs/:id/overview", "admin"),
    ("jobs-detail-pipeline_u.png", "岗位详情 — Pipeline 管道视图（9阶段+卡点标注）", "GET /api/jobs/:id/pipeline", "admin"),
    ("jobs-detail-candidate-mapping_u.png", "岗位详情 — 候选人分布+状态+卡点标签", "GET /api/jobs/:id/overview", "admin"),
]

def set_cell_shading(cell, color):
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color,
    })
    shading_elm.append(shading)

def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '2B579A')
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r % 2 == 0:
                set_cell_shading(cell, 'F2F6FC')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# Title
title = doc.add_heading('理然智能招聘 AI 看板 — Phase 8.1 Jobs 模块 自检报告', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f'生成日期：{datetime.now().strftime("%Y-%m-%d %H:%M")}')
doc.add_paragraph('分支：agent/workbuddy/phase-7')
doc.add_paragraph('阶段：Phase 8.1 — Jobs 模块 SaaS 闭环')
doc.add_paragraph('Mock：否 — 全部截图来自真实 API 渲染')
doc.add_paragraph()

# 一、构建验证
doc.add_heading('一、构建验证', level=1)
add_styled_table(doc,
    ['检查项', '结果', '说明'],
    [
        ['Type Check', '✅ PASS', '0 errors'],
        ['Lint', '✅ PASS', '0 errors 0 warnings'],
        ['Build', '✅ PASS', '全部路由编译，新增 /api/jobs/:id/pipeline, /api/jobs/:id/overview'],
    ]
)

# 二、Jobs 模块 SaaS 闭环
doc.add_heading('二、Jobs 模块 SaaS 闭环实现', level=1)
add_styled_table(doc,
    ['Taskbook 要求', '实现状态', '交付物'],
    [
        ['Job Overview（人数/进度/成功率/风险）', '✅', 'GET /api/jobs/:id/overview + JobOverview 组件'],
        ['Pipeline View（阶段转化+卡点）', '✅', 'GET /api/jobs/:id/pipeline + JobPipelineView 组件'],
        ['Bottleneck Intelligence（卡点原因=rule解释）', '✅', 'Rule Engine Rule 7（14天停留检测）+ Pipeline 卡点标注'],
        ['Candidate Mapping（候选人分布+状态+风险）', '✅', 'Overview API 含 candidateMapping 字段'],
        ['Action Link（规则触发Action绑定Job+Candidate）', '✅', 'Pipeline/Overview 均展示 openActions + 跳转入口'],
        ['Rule Engine 接入', '✅', '新增 Rule 7（流程卡点检测）+ 已有 6 条规则'],
    ]
)

# 三、新增/修改文件清单
doc.add_heading('三、文件清单', level=1)
add_styled_table(doc,
    ['文件', '说明'],
    [
        ['app/api/jobs/[id]/pipeline/route.ts', 'Pipeline API：9阶段人数+转化率+卡点检测'],
        ['app/api/jobs/[id]/overview/route.ts', 'Overview API：编制进度+风险+候选人分布+Action'],
        ['components/domain/jobs/JobPipelineView.tsx', '水平阶段管道视图组件'],
        ['components/domain/jobs/JobOverview.tsx', '岗位总览组件（KPI+风险+候选人分布）'],
        ['app/jobs/page.tsx', '重写：列表→详情双视图，整合 Overview+Pipeline'],
        ['server/services/action/action-rule-service.ts', '新增 Rule 7：流程卡点检测（14天停留）'],
    ]
)

# 四、截图证据
doc.add_heading('四、截图证据（共 4 张）', level=1)
p = doc.add_paragraph()
p.add_run('以下 4 张截图全部来自真实 API 渲染，无 Mock 数据。').bold = True

for i, (filename, description, api_info, role) in enumerate(screenshots, 1):
    doc.add_heading(f'{i}. {description}', level=2)
    add_styled_table(doc,
        ['属性', '值'],
        [
            ['文件名', filename],
            ['描述', description],
            ['数据来源', api_info],
            ['角色', role],
            ['Mock', '否'],
        ],
        col_widths=[3, 13]
    )
    filepath = os.path.join(SCREENSHOT_DIR, filename)
    if os.path.exists(filepath):
        doc.add_picture(filepath, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f'⚠️ 截图文件未找到: {filepath}')
    doc.add_paragraph()

# 五、验收红线
doc.add_heading('五、验收红线确认', level=1)
add_styled_table(doc,
    ['#', '红线项', '状态'],
    [
        ['1', 'Job Overview 可访问', '✅'],
        ['2', 'Pipeline View 可访问', '✅'],
        ['3', 'Bottleneck 检测正常', '✅'],
        ['4', 'Candidate Mapping 正常', '✅'],
        ['5', 'Action Link 绑定 Job+Candidate', '✅'],
        ['6', 'Rule Engine 驱动', '✅'],
        ['7', '无 mock/demo 数据', '✅'],
        ['8', '无 AI 决策', '✅'],
        ['9', 'typecheck/lint/build 通过', '✅'],
        ['10', '截图完成（4 张）', '✅'],
        ['11', '未破坏 Action Center', '✅'],
        ['12', '未做 Candidates 模块', '✅'],
    ]
)

# 六、最终结论
doc.add_heading('六、最终结论', level=1)
add_styled_table(doc,
    ['项目', '结论'],
    [
        ['Phase 8.1 Jobs 模块 SaaS 闭环是否完成', '是'],
        ['Job Overview 是否实现', '是'],
        ['Pipeline View 是否实现', '是'],
        ['Bottleneck Intelligence 是否实现', '是'],
        ['Candidate Mapping 是否实现', '是'],
        ['Action Link 是否实现', '是'],
        ['Rule Engine 是否接入', '是'],
        ['是否提前做 Candidates', '否'],
        ['typecheck/lint/build 是否通过', '是'],
        ['截图是否完成', '是（4 张）'],
        ['是否使用 mock 数据', '否'],
        ['是否存在假 AI', '否'],
        ['需要外部确认', 'ChatGPT 最终验收'],
    ]
)

doc.save(OUTPUT)
print(f"✅ Word report: {OUTPUT}")
print(f"   Screenshots: {len(screenshots)}")
