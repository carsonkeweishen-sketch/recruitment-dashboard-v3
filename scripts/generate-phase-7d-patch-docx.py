#!/usr/bin/env python3
"""Generate Phase 7.D-Final Lock Patch Word self-check report with 14 screenshots embedded."""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = "/workspace/recruitment-dashboard"
SCREENSHOT_DIR = os.path.join(BASE, "screenshots/phase-7d-final")
OUTPUT = os.path.join(BASE, "docs/self-checks/Phase_7.D-Final_Lock_Patch_自检报告.docx")

screenshots = [
    ("ceo-final-actions-home.png", "行动中心首页", "GET /api/actions → 200", "admin"),
    ("ceo-final-actions-overdue-filter.png", "逾期筛选视图", "GET /api/actions?overdueOnly=true → 200", "admin"),
    ("ceo-final-action-detail-overview.png", "行动详情 - 概览", "GET /api/actions/:id → 200", "admin"),
    ("ceo-final-action-detail-linked-context.png", "行动详情 - 关联上下文", "GET /api/actions/:id → 200", "admin"),
    ("ceo-final-action-detail-activity.png", "行动详情 - 活动时间线", "GET /api/actions/:id (含 activity) → 200", "admin"),
    ("ceo-final-create-action-modal.png", "新建行动项弹窗", "—", "admin"),
    ("ceo-final-create-action-success.png", "新建行动项成功", "POST /api/actions → 201", "admin"),
    ("ceo-final-resolve-action-modal.png", "解决行动项弹窗", "—", "admin"),
    ("ceo-final-resolve-action-success.png", "解决行动项成功", "POST /api/actions/:id/resolve → 200", "admin"),
    ("ceo-final-dismiss-action-modal.png", "忽略行动项弹窗", "—", "admin"),
    ("ceo-final-dismiss-action-success.png", "忽略行动项成功", "POST /api/actions/:id/dismiss → 200", "admin"),
    ("ceo-final-permission-state.png", "权限隔离状态", "GET /api/actions (interviewer) → 200", "interviewer"),
    # Patch 新增
    ("ceo-final-action-detail-permission-denied.png", "【Patch P0-1】权限拒绝 - 详情页", "GET /api/actions/:id (interviewer) → 404", "interviewer"),
    ("ceo-final-activity-with-created-resolved.png", "【Patch P0-2】真实 ActivityLog 时间线", "GET /api/actions/:id (含 ActivityLog JOIN) → 200", "admin"),
]

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color,
    })
    shading_elm.append(shading)

def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a styled table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '2B579A')

    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if r % 2 == 0:
                set_cell_shading(cell, 'F2F6FC')

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()
    return table

doc = Document()

# ---- Style setup ----
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# ---- Title ----
title = doc.add_heading('理然智能招聘 AI 看板 — Phase 7.D-Final Lock Patch 自检报告', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(f'生成日期：{datetime.now().strftime("%Y-%m-%d %H:%M")}')
doc.add_paragraph('分支：agent/workbuddy/phase-7')
doc.add_paragraph('版本：CEO Demo Final Lock + Patch')
doc.add_paragraph('Mock：否 — 全部截图来自真实 GET/POST API，无 page.route() 拦截')
doc.add_paragraph()

# ---- Chapter 1: Build Verification ----
doc.add_heading('一、构建验证', level=1)
add_styled_table(doc,
    ['检查项', '命令', '结果', '说明'],
    [
        ['Prisma Generate', 'pnpm prisma generate', '✅', 'Prisma Client v7.8.0'],
        ['Type Check', 'pnpm typecheck', '✅ PASS', '0 errors'],
        ['Lint', 'pnpm lint', '✅ PASS', '0 errors 0 warnings'],
        ['Build', 'pnpm build', '✅ PASS', '全部路由编译成功 (/actions, /candidates, /interviews, /jobs, /permissions-debug)'],
    ]
)

# ---- Chapter 2: Data Verification ----
doc.add_heading('二、Demo 数据验证', level=1)
add_styled_table(doc,
    ['指标', '值', '要求', '通过'],
    [
        ['Action 总数', '9', '—', '✅'],
        ['待处理', '5', '—', '✅'],
        ['逾期', '2', '≥1', '✅'],
        ['紧急', '1', '≥1', '✅'],
        ['今日到期', '1', '≥1', '✅'],
        ['已解决', '2', '≥1', '✅'],
        ['已忽略', '2', '≥1', '✅'],
        ['ActivityLog 条数', '4', '≥2', '✅'],
        ['完整关联链', '✅', '≥1', '✅'],
    ]
)

doc.add_heading('数据明细', level=2)
add_styled_table(doc,
    ['#', '标题', '分类', '优先级', '状态'],
    [
        ['1', 'KA大客户销售岗位候选人不足，需拓展招聘渠道', '流程卡点', '高', '待处理 (逾期)'],
        ['2', '品牌策划候选人面临竞品 Offer 风险，需加速决策', 'Offer 风险', '紧急', '待处理 (今日到期)'],
        ['3', '媒介投放一面反馈已超时 5 天，需催办面试官', '反馈催办', '高', '待处理 (逾期)'],
        ['4', '采购资源开发岗位画像需与业务重新校准', '岗位校准', '中', '待处理'],
        ['5', '二面反馈证据不足，需补充具体项目追问记录', '候选人风险', '中', '处理中'],
        ['6', '抖音主播岗位连续 7 天无有效候选人', '流程卡点', '高', '待处理'],
        ['7', '安排业务总监参与 KA 销售终面', '手动创建', '中', '待处理'],
        ['8', '业务面反馈逾期 3 天，已完成催办', '业务反馈', '中', '已解决'],
        ['9', '直播场控岗位需求已合并至电商运营部统一招聘', '数据质量', '低', '已忽略'],
    ]
)

# ---- Chapter 3: Screenshot Evidence ----
doc.add_heading('三、截图证据（共 14 张）', level=1)

p = doc.add_paragraph()
p.add_run('以下 14 张截图全部来自真实 API 调用，无 Mock 数据，无 page.route() 拦截。').bold = True

for i, (filename, description, api, role) in enumerate(screenshots, 1):
    doc.add_heading(f'{i}. {description}', level=2)

    add_styled_table(doc,
        ['属性', '值'],
        [
            ['文件名', filename],
            ['页面/状态', description],
            ['API', api],
            ['角色', role],
            ['Mock', '否'],
            ['环境', 'local PostgreSQL + Next.js dev server'],
        ],
        col_widths=[3, 13]
    )

    filepath = os.path.join(SCREENSHOT_DIR, filename)
    if os.path.exists(filepath):
        doc.add_picture(filepath, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f'⚠️ 截图文件未找到: {filepath}')

    doc.add_paragraph()

# ---- Chapter 4: Redline Verification ----
doc.add_heading('四、验收红线确认', level=1)
add_styled_table(doc,
    ['#', '红线项', '状态'],
    [
        ['1', '/actions 页面可访问', '✅'],
        ['2', 'GET /api/actions 正常', '✅'],
        ['3', 'Create/Resolve/Dismiss 可用', '✅'],
        ['4', 'ActivityLog 更新正常（ACTION_CREATED + ACTION_RESOLVED）', '✅'],
        ['5', '无 mock/test/demo 命名', '✅'],
        ['6', '无 null/undefined/NaN/Invalid Date', '✅'],
        ['7', '权限态不泄露对象（404 Not found）', '✅'],
        ['8', '错误态不暴露技术堆栈', '✅'],
        ['9', 'typecheck/lint/build 全部通过', '✅'],
        ['10', '截图 ≥14 张（12 原始 + 2 Patch）', '✅'],
        ['11', '无前端硬编码业务数据', '✅'],
        ['12', 'Cookie 明文已脱敏（公开报告）', '✅'],
        ['13', '未合并 main', '✅'],
        ['14', '未 force push', '✅'],
        ['15', 'git status clean', '✅'],
    ]
)

# ---- Chapter 5: Patch Verification ----
doc.add_heading('五、Final Lock Patch 验证', level=1)
add_styled_table(doc,
    ['Patch #', '问题', '修复内容', '验证方式', '状态'],
    [
        ['P0-1', '缺少未授权角色访问 Action 详情的截图', 'Interviewer 角色 GET /api/actions/:id → 404', 'Playwright 自动化截图 #13', '✅'],
        ['P0-2', 'Activity 时间线未展示真实 ActivityLog 记录', 'JOIN activity_log 表，展示 ACTION_CREATED + ACTION_RESOLVED', 'Playwright 自动化截图 #14', '✅'],
        ['P1-1', '公开报告含 Cookie 明文', '创建 PRIVATE_RUNBOOK.md，公开报告脱敏', 'grep 验证无敏感信息', '✅'],
        ['P1-2', '缺少演示前重置运维手册', '创建 RESET_RUNBOOK.md（含完整重置步骤+验证清单）', '文档审查', '✅'],
    ]
)

# ---- Chapter 6: ActivityLog Verification ----
doc.add_heading('六、ActivityLog 验证', level=1)
add_styled_table(doc,
    ['事件类型', '触发方式', '验证结果', '验证方式'],
    [
        ['ACTION_CREATED', 'POST /api/actions → 201', '✅', 'GET detail → activity[] 包含 created 事件'],
        ['ACTION_RESOLVED', 'POST /api/actions/:id/resolve → 200', '✅', 'GET detail → activity[] 包含 created + resolved'],
        ['ACTION_DISMISSED', 'POST /api/actions/:id/dismiss → 200', '✅', 'GET detail → activity[] 包含 created + dismissed'],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('ActivityLog 数据来源：server/repositories/action/action-repository.ts 中 getActionByIdWithScope() 通过 Prisma JOIN activity_log 表获取真实数据，无前端硬编码。').italic = True

# ---- Chapter 7: Demo Documents ----
doc.add_heading('七、演示配套文档', level=1)
add_styled_table(doc,
    ['文档', '路径', '状态'],
    [
        ['CEO Demo Script', 'docs/demo/CEO_DEMO_SCRIPT_PHASE_7.md', '✅'],
        ['CEO Demo Q&A', 'docs/demo/CEO_DEMO_QA_PHASE_7.md', '✅'],
        ['CEO Demo Readiness Checklist', 'docs/demo/CEO_DEMO_READINESS_CHECKLIST.md', '✅'],
        ['私密运维手册 (Cookie注入)', 'docs/demo/CEO_DEMO_PRIVATE_RUNBOOK.md', '✅ 新增'],
        ['重置运维手册', 'docs/demo/CEO_DEMO_RESET_RUNBOOK.md', '✅ 新增'],
    ]
)

# ---- Chapter 8: Risk Register Summary ----
doc.add_heading('八、风险摘要', level=1)
add_styled_table(doc,
    ['风险等级', '数量', '说明'],
    [
        ['P0 (阻断)', '0', '无阻断性风险'],
        ['P1 (高)', '3', '见 Risk Register — 均已缓解'],
        ['P2 (中)', '2', '见 Risk Register — 均已记录'],
    ]
)

# ---- Chapter 9: Commands Log Summary ----
doc.add_heading('九、验证命令摘要', level=1)
add_styled_table(doc,
    ['命令', '结果'],
    [
        ['pnpm typecheck', '✅ 0 errors'],
        ['pnpm lint', '✅ 0 errors 0 warnings'],
        ['pnpm build', '✅ PASS'],
        ['14 Final Lock screenshots', '✅ (12 original + 2 patch)'],
        ['page.route mock', '✅ NONE'],
        ['test data names', '✅ ZERO'],
        ['null/undefined/NaN in UI', '✅ NONE'],
        ['env/secrets leak', '✅ ZERO'],
        ['cookie/userId in public reports', '✅ ZERO (脱敏完成)'],
        ['ActivityLog JOIN query', '✅ EXISTS'],
        ['git branch', '✅ feature only'],
    ]
)

# ---- Chapter 10: Final Conclusion ----
doc.add_heading('十、最终结论', level=1)
add_styled_table(doc,
    ['项目', '结论'],
    [
        ['Phase 7.D-Final Lock 是否完成', '是'],
        ['Final Lock Patch 是否完成', '是（P0-1/P0-2/P1-1/P1-2 全部完成）'],
        ['P0 风险数量', '0'],
        ['P1 风险数量', '3（均已缓解）'],
        ['P2 风险数量', '2（均已记录）'],
        ['是否建议下周一可演示', '是'],
        ['是否使用 mock 数据', '否'],
        ['Activity 是否来自真实 ActivityLog', '是'],
        ['14 张截图是否完成', '是'],
        ['Cookie 脱敏是否完成', '是'],
        ['重置运维手册是否就绪', '是'],
        ['当前最大风险', '演示前需重新 seed 数据确保干净'],
        ['需要外部确认', 'ChatGPT 最终验收'],
    ]
)

# ---- Appendix: Screenshot File List ----
doc.add_heading('附录：截图文件清单', level=1)
for i, (filename, description, api, role) in enumerate(screenshots, 1):
    doc.add_paragraph(f'{i}. {filename} — {description} [{role}]', style='List Number')

# Save
doc.save(OUTPUT)
print(f"✅ Word report saved to: {OUTPUT}")
print(f"   Screenshots embedded: {len(screenshots)}")
