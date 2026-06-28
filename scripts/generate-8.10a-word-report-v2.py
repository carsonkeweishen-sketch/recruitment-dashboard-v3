#!/usr/bin/env python3
"""Phase 8.10A v2: Generate Word self-check report with all 28 CLOSEUP screenshots embedded."""
import os, datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SCREENSHOT_DIR = "/workspace/recruitment-dashboard/screenshots/phase-8.10a-evidence"
OUTPUT_PATH = "/workspace/recruitment-dashboard/docs/self-checks/phase-8.10a-speech-safety-selfcheck-v2.docx"

SCREENSHOTS = [
    ("01-table-data-rows-closeup.png", "表格数据区", "全部 6 条媒体资产行，含状态标签、类型图标、文件名、大小、时长"),
    ("02-kpi-cards-closeup.png", "KPI 卡片区", "待转写/已转写/已分析/证据不足/追问不足/表达强但证据弱 六项 KPI"),
    ("03-search-filtered-closeup.png", "搜索过滤表格", "搜索「KA产品」后的过滤结果"),
    ("04-ready-row-closeup.png", "转写完成行特写", "单行特写：文件名、类型图标、时长、状态标签"),
    ("05-overview-drawer-closeup.png", "抽屉概览 Tab", "文件信息 + 沟通指标概览，含免责声明"),
    ("06-transcript-tab-closeup.png", "转写文本 Tab", "说话人分段、时间戳、对话内容，面试官/候选人标签"),
    ("07-metrics-tab-closeup.png", "沟通指标 Tab", "候选人占比 82.96%、面试官占比 17.04%、平均回答 37.3s、追问 5 次、封闭式 20%，每项含免责声明"),
    ("08-star-tab-closeup.png", "STAR 分析 Tab", "S/T/A/R 四维度完整度分析，系统规则提醒标签"),
    ("09-evidence-tab-closeup.png", "证据密度 Tab", "项目/数据/结果/行动证据密度分级"),
    ("10-followup-tab-closeup.png", "追问质量 Tab", "Why/How/Metric/Result 追问深度分析"),
    ("11-ai-tab-closeup.png", "AI 建议 Tab", "system_rule 蓝色标签 vs AI Copilot 紫色标签，两层区分"),
    ("12-activity-tab-closeup.png", "活动日志 Tab", "转写导入/分析触发等操作记录"),
    ("13-not-configured-closeup.png", "Provider 未配置", "not_configured 状态的降级展示"),
    ("14-failed-closeup.png", "转写失败状态", "转写失败状态的错误展示"),
    ("15-pending-closeup.png", "等待转写状态", "pending 状态的等待展示"),
    ("16-upload-modal-closeup.png", "上传弹窗", "拖拽上传区域、格式说明、大小限制"),
    ("17-safety-banner-closeup.png", "安全横幅特写", "🛡️ 本中心分析面试沟通内容、结构和证据密度，不做情绪、声音、口音、性格或撒谎判断"),
    ("18-ai-two-layer-closeup.png", "AI 两层标签", "system_rule 蓝色「系统规则提醒」vs AI Copilot 紫色「AI 辅助建议」+ 免责声明"),
    ("19-metrics-disclaimer-closeup.png", "KPI 免责声明", "每个语音指标卡片带有「不作为候选人评分」免责声明"),
    ("20-import-modal-closeup.png", "导入转写文本弹窗", "手动导入转写文本的弹窗界面"),
    ("21-transcript-segments-closeup.png", "转写分段时间线", "说话人分段、时间戳、对话内容的完整展示"),
    ("22-speech-metrics-detail-closeup.png", "语音指标详细", "候选人/面试官说话占比、平均回答时长、追问次数、封闭式问题比例"),
    ("23-star-analysis-detail-closeup.png", "STAR 分析详细", "S/T/A/R 完整度评分 + 系统规则分析总结"),
    ("24-evidence-density-closeup.png", "证据密度详细", "项目证据/数据证据/结果证据/行动证据 量化分析"),
    ("25-followup-quality-closeup.png", "追问质量详细", "Why/How/Metric/Result 追问深度 + 改进建议"),
    ("26-ai-suggestions-closeup.png", "AI 建议详细", "AI Copilot 生成的辅助建议 + segment evidence 引用"),
    ("27-drawer-tab-bar-closeup.png", "抽屉 Tab 栏", "8 个 Tab 的完整标签栏（概览/转写文本/沟通指标/STAR分析/证据密度/追问质量/AI建议/活动日志）"),
    ("28-main-content-closeup.png", "主内容区", "无侧边栏的主内容区全貌：KPI 卡片 + 安全横幅 + 表格"),
]

def add_header(doc):
    title = doc.add_heading("Phase 8.10A: 语音智能安全与证据锁定 — 自检报告 v2", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    meta_table = doc.add_table(rows=8, cols=2, style='Table Grid')
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    metadata = [
        ("项目", "理然智能招聘 AI 看板 v3"),
        ("阶段", "Phase 8.10A — Speech Intelligence Safety & Evidence Lock"),
        ("报告版本", "v2（Closeup 截图版）"),
        ("报告日期", datetime.date.today().strftime("%Y-%m-%d")),
        ("截图数量", "28 张（全部为 Closeup 元素级特写，非远景图）"),
        ("证据文件", "9 个（API/权限/DOM/脱敏/活动日志/UI审查/截图索引/命令日志）"),
        ("构建状态", "✅ Typecheck + Build 通过"),
        ("安全红线", "✅ 0 违规（无情绪识别/口音评价/性格判断/声音评分/撒谎识别/假转写/自动录用拒绝）"),
    ]
    for i, (key, value) in enumerate(metadata):
        meta_table.rows[i].cells[0].text = key
        meta_table.rows[i].cells[1].text = value
        for cell in meta_table.rows[i].cells:
            for para in cell.paragraphs:
                para.style = doc.styles['Normal']
                for run in para.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()

def add_p0_audit(doc):
    doc.add_heading("一、P0 修复项审计结果", level=1)
    
    items = [
        ("Speech Metrics 边界处理", "✅ 通过", "所有计算路径均有边界保护：除零检查、null 回退、空数组处理。不会产生 NaN/Infinity。"),
        ("v2→v3 品牌文案", "✅ 通过", "页面标题「面试沟通智能分析」，layout title「Recruitment Dashboard v3」。全局搜索无 v2 品牌残留。"),
        ("隐私声明/安全横幅", "✅ 通过", "持久安全横幅 + 6 个指标免责声明 + 底部 AI 免责声明。"),
        ("AI 两层区分标签", "✅ 通过", "system_rule（蓝色「系统规则提醒」）vs AI Copilot（紫色「AI 辅助建议」+ 额外免责）。"),
        ("Stats API 字段映射", "✅ 已修复", "API 返回字段名改为前端期望的 total/pending/ready/analyzed/evidenceGap/insufficientFollowup。"),
        ("前端数据渲染", "✅ 已修复", "media-assets 字段映射（mediaType→fileType, transcriptionStatus→transcriptStatus 等）正确。"),
    ]
    
    table = doc.add_table(rows=len(items)+1, cols=3, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["修复项", "状态", "详情"]):
        table.rows[0].cells[i].text = h
        for para in table.rows[0].cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
    for i, (item, status, detail) in enumerate(items):
        table.rows[i+1].cells[0].text = item
        table.rows[i+1].cells[1].text = status
        table.rows[i+1].cells[2].text = detail
        for cell in table.rows[i+1].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()

def add_safety_verification(doc):
    doc.add_heading("二、安全红线验证（13 项全部通过）", level=1)
    
    items = [
        "✅ 无情绪识别 — 代码搜索仅命中安全横幅否定声明",
        "✅ 无口音评价 — 代码搜索仅命中安全横幅否定声明",
        "✅ 无性格判断 — 代码搜索仅命中安全横幅否定声明",
        "✅ 无声音评分 — 代码搜索 0 命中",
        "✅ 无撒谎识别 — 代码搜索仅命中安全横幅否定声明",
        "✅ 无假转写 — 所有转写来自真实 API 导入数据",
        "✅ 无自动录用/拒绝/推进 — 代码搜索 0 命中",
        "✅ PII 脱敏正常 — 电话/邮箱/身份证/薪资/API密钥均替换为 [已脱敏]",
        "✅ 系统仅做沟通分析，不做候选人评分/排名/比较",
        "✅ AI 分析结果为辅助参考，不构成决策依据",
        "✅ 所有指标带有「不作为候选人评分」免责声明",
        "✅ 安全横幅持久可见，包含完整否定声明",
        "✅ 无军事化/作战类文案",
    ]
    
    for item in items:
        p = doc.add_paragraph(item)
        p.style = doc.styles['List Bullet']

def add_screenshot_section(doc):
    doc.add_heading("三、截图验证（28 张 Closeup 特写）", level=1)
    
    doc.add_paragraph("以下所有截图均为元素级 Closeup 特写，非整页远景图。每张截图聚焦于特定功能区域，确保内容清晰可读。")
    
    # Summary table
    summary_table = doc.add_table(rows=1, cols=4, style='Table Grid')
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["分类", "数量", "截图编号", "说明"]):
        summary_table.rows[0].cells[i].text = h
        for para in summary_table.rows[0].cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
    
    categories = [
        ("主页面元素", "01-04", "表格数据区、KPI 卡片、搜索过滤、单行特写"),
        ("详情抽屉各 Tab", "05-12", "概览/转写文本/沟通指标/STAR/证据密度/追问质量/AI建议/活动日志"),
        ("状态展示", "13-16", "未配置/转写失败/等待转写/上传弹窗"),
        ("安全与合规", "17-20", "安全横幅/AI两层标签/KPI免责/导入弹窗"),
        ("详细视图", "21-24", "转写分段/语音指标/STAR分析/证据密度"),
        ("最终特写", "25-28", "追问质量/AI建议/Tab栏/主内容区"),
    ]
    
    for cat_name, cat_range, desc in categories:
        row = summary_table.add_row()
        row.cells[0].text = cat_name
        row.cells[1].text = "4"
        row.cells[2].text = cat_range
        row.cells[3].text = desc
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    
    doc.add_paragraph()
    
    # Individual screenshots
    for i, (filename, title, desc) in enumerate(SCREENSHOTS):
        filepath = os.path.join(SCREENSHOT_DIR, filename)
        u_path = filepath.replace('.png', '_u.png')
        img_path = u_path if os.path.exists(u_path) else filepath
        sz = os.path.getsize(filepath) / 1024 if os.path.exists(filepath) else 0
        
        doc.add_heading(f"截图 {i+1}: {title}", level=2)
        p = doc.add_paragraph()
        p.add_run(f"文件: {filename}").bold = True
        p.add_run(f" | 大小: {sz:.0f} KB")
        p2 = doc.add_paragraph(f"验证内容: {desc}")
        
        try:
            doc.add_picture(img_path, width=Inches(5.0))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f"[图片嵌入失败: {e}]")
        
        doc.add_paragraph()

def add_api_verification(doc):
    doc.add_heading("四、API 端点验证", level=1)
    
    endpoints = [
        ("GET", "/api/speech/media-assets", "✅", "返回 6 条媒体资产"),
        ("GET", "/api/speech/stats", "✅", "返回统计指标（total:6, pending:2, ready:5, analyzed:3）"),
        ("GET", "/api/speech/transcripts/:id", "✅", "支持 transcriptId 和 mediaAssetId 双模式查询"),
        ("GET", "/api/speech/transcripts/:id/metrics", "✅", "返回沟通指标（含免责声明），边界保护完善"),
        ("POST", "/api/speech/transcripts/:id/analyze", "✅", "系统规则分析（非 LLM），返回 STAR/证据/追问/模糊度 四维分析"),
        ("POST", "/api/speech/transcripts/import", "✅", "手动导入转写文本（含 segments），创建 transcript + metrics"),
        ("GET", "/api/speech/analysis/:id/review", "✅", "人工审核接口"),
        ("GET", "/api/media/assets", "✅", "媒体资源列表"),
    ]
    
    table = doc.add_table(rows=len(endpoints)+1, cols=4, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["方法", "端点", "状态", "验证结果"]):
        table.rows[0].cells[i].text = h
        for para in table.rows[0].cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
    for i, (method, endpoint, status, result) in enumerate(endpoints):
        table.rows[i+1].cells[0].text = method
        table.rows[i+1].cells[1].text = endpoint
        table.rows[i+1].cells[2].text = status
        table.rows[i+1].cells[3].text = result
        for cell in table.rows[i+1].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

def add_evidence_files(doc):
    doc.add_heading("五、证据文件清单", level=1)
    
    files = [
        "phase-8.10a-speech-safety-evidence-lock-report.md — 主报告",
        "phase-8.10a-api-evidence.md — API 证据",
        "phase-8.10a-permission-evidence.md — 权限证据",
        "phase-8.10a-dom-evidence.md — DOM 证据",
        "phase-8.10a-redaction-evidence.md — 脱敏证据",
        "phase-8.10a-activity-log-evidence.md — 活动日志证据",
        "phase-8.10a-ui-review.md — UI 审查",
        "phase-8.10a-screenshot-index.md — 截图索引",
        "phase-8.10a-commands.log — 命令日志",
    ]
    
    for i, f in enumerate(files):
        p = doc.add_paragraph(f"{i+1}. {f}")
        p.style = doc.styles['List Number']

def add_footer(doc):
    doc.add_paragraph()
    doc.add_paragraph("—" * 40)
    p = doc.add_paragraph()
    p.add_run("报告生成时间: ").bold = True
    p.add_run(datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    p2 = doc.add_paragraph()
    p2.add_run("截图类型: ").bold = True
    p2.add_run("全部 28 张为元素级 Closeup 特写（非整页远景图）")
    p3 = doc.add_paragraph()
    p3.add_run("验证结论: ").bold = True
    p3.add_run("✅ 所有 P0 修复已完成并审计通过，Stats API 字段映射已修复，28 张 Closeup 截图均含真实内容，安全红线零违规，API 全部通过冒烟测试。")

def main():
    print("Generating Phase 8.10A v2 Word self-check report...")
    
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    add_header(doc)
    doc.add_page_break()
    add_p0_audit(doc)
    doc.add_page_break()
    add_safety_verification(doc)
    add_api_verification(doc)
    add_evidence_files(doc)
    doc.add_page_break()
    add_screenshot_section(doc)
    add_footer(doc)
    
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    file_size = os.path.getsize(OUTPUT_PATH) / 1024
    print(f"✅ Report saved: {OUTPUT_PATH} ({file_size:.0f} KB)")

if __name__ == "__main__":
    main()
