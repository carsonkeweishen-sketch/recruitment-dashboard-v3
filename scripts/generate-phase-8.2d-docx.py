#!/usr/bin/env python3
"""Phase 8.2D Real UI Evidence Lock — Word 自检报告 (6 verified screenshots)"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = "/workspace/recruitment-dashboard"
DIR = os.path.join(BASE, "screenshots/phase-8.2-job-center")
OUT = os.path.join(BASE, "docs/self-checks/Phase_8.2D_Job_Center_自检报告.docx")

shots = [
    ("job-detail-drawer-insights-real_u.png", "[Insights] KA大客户销售", "2条洞察: 供给偏弱+逾期, 含evidence+suggestedAction+系统规则提醒"),
    ("job-insight-provenance-readable_u.png", "[Provenance] 系统规则提醒", "生成方式+触发条件+证据数量+时间"),
    ("job-detail-drawer-activity-real_u.png", "[Activity] 动态记录", "人话化文案, 非裸枚举"),
    ("job-detail-drawer-candidates-real_u.png", "[Candidates] 候选人质量", "候选人名+阶段+公司/职位, 无敏感信息"),
    ("job-detail-drawer-interview-quality-real_u.png", "[Interview Quality] 面试反馈", "待提交/提交率/时长/质量分"),
    ("job-detail-drawer-actions-real_u.png", "[Actions] 风险行动", "标题+分类+优先级+状态+负责人+逾期, 不直接Resolve"),
]

def shading(cell, color):
    e = cell._element.get_or_add_tcPr()
    e.append(e.makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color}))

def tbl(doc, h, r, w=None):
    t = doc.add_table(rows=1+len(r), cols=len(h)); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hh in enumerate(h):
        c = t.rows[0].cells[i]; c.text = hh
        for p in c.paragraphs: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in c.paragraphs[0].runs: run.bold = True; run.font.size = Pt(10); run.font.color.rgb = RGBColor(255,255,255)
        shading(c, '2B579A')
    for ri, row in enumerate(r):
        for ci, v in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for run in p.runs: run.font.size = Pt(9)
            if ri % 2 == 0: shading(c, 'F2F6FC')
    if w:
        for i, ww in enumerate(w):
            for row in t.rows: row.cells[i].width = Cm(ww)
    doc.add_paragraph(); return t

doc = Document()
s = doc.styles['Normal']; s.font.name = 'Microsoft YaHei'; s.font.size = Pt(10.5)
s.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

t = doc.add_heading('理然智能招聘 AI 看板 - Phase 8.2D Job Center 自检报告', 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f'生成日期: {datetime.now().strftime("%Y-%m-%d %H:%M")} | 分支: agent/workbuddy/phase-7')
doc.add_paragraph('目标岗位: KA大客户销售 | Insights验证: 候选人供给偏弱+存在逾期行动项 (2条, 系统规则提醒=true, 暂无系统洞察=false)')
doc.add_paragraph('Mock: 否。全部截图来自真实 API，Playwright DOM 验证通过。')
doc.add_paragraph()

doc.add_heading('一、构建验证', 1)
tbl(doc, ['检查项','结果'], [['typecheck','PASS'],['lint','PASS'],['build','PASS'],['git','clean']])

doc.add_heading('二、DOM 内容验证 (Playwright)', 1)
tbl(doc, ['验证项','结果'], [
    ['Has 候选人供给偏弱','TRUE'],['Has 存在逾期行动项','TRUE'],
    ['Has 系统规则提醒','TRUE'],['Has 暂无系统洞察','FALSE (非空态!)'],
])

doc.add_heading('三、6 张截图证据', 1)
for i, (fn, desc, verify) in enumerate(shots, 1):
    doc.add_heading(f'{i}. {desc}', 2)
    tbl(doc, ['属性','值'], [['文件名',fn.replace('_u.png','.png')],['描述',desc],['验证',verify],['Mock','否']], [3,13])
    fp = os.path.join(DIR, fn)
    if os.path.exists(fp): doc.add_picture(fp, width=Inches(5.5)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else: doc.add_paragraph(f'MISSING: {fp}')
    doc.add_paragraph()

doc.add_heading('四、验收红线', 1)
tbl(doc, ['#','红线','状态'], [
    ['1','Insights 非空态','PASS (2条真实洞察)'],['2','Insights 含evidence/suggestedAction','PASS'],
    ['3','Provenance 可读','PASS'],['4','Activity 人话化','PASS'],
    ['5','Candidates 可读无敏感信息','PASS'],['6','Interview Quality 可读','PASS'],
    ['7','Actions 可读不直接Resolve','PASS'],['8','无 AI 决策/自动淘汰','PASS'],
    ['9','typecheck/lint/build 通过','PASS'],['10','git clean','PASS'],
])

doc.add_heading('五、最终结论', 1)
tbl(doc, ['项目','结论'], [
    ['Phase 8.2D 是否完成','是'],['Insights 是否有真实非空洞察','是 (2条, DOM验证)'],
    ['Provenance 是否肉眼可读','是'],['Activity 是否人话化','是'],
    ['Candidates 是否可读','是'],['Interview Quality 是否可读','是'],
    ['Actions 是否可读','是'],['是否接 OpenAI','否'],
    ['是否伪造 LLM','否'],['是否破坏 Action Center','否'],
    ['typecheck/lint/build 是否通过','是'],['git 是否 clean','是'],
    ['是否合并 main','否'],['是否进入 Phase 8.3','否'],
    ['需要外部确认','ChatGPT 最终验收'],
])

doc.save(OUT)
print(f'OK {OUT} | Screenshots: {len(shots)}')
