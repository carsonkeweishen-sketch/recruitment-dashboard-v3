#!/usr/bin/env python3
"""Phase 8.12 + 8.12B: Combined self-check report for GPT review"""
import os, datetime
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Screenshot directories
DIR_812 = "/workspace/recruitment-dashboard/screenshots/phase-8.12-v2"
DIR_812B = "/workspace/recruitment-dashboard/screenshots/phase-8.12b"
OUTPUT = "/workspace/recruitment-dashboard/docs/self-checks/phase-8.12-combined-selfcheck.docx"

# All screenshots with metadata
SCREENSHOTS = []

# Phase 8.12 screenshots (52)
for i in range(1, 53):
    fn = f"{i:02d}-"
    # Find actual file
    for f in os.listdir(DIR_812):
        if f.startswith(fn) and f.endswith('.png') and '_u' not in f:
            SCREENSHOTS.append((os.path.join(DIR_812, f), f, "Phase 8.12", f"8.12 截图 #{i}"))
            break

# Phase 8.12B screenshots (10)
for f in sorted(os.listdir(DIR_812B)):
    if f.endswith('.png') and '_u' not in f:
        SCREENSHOTS.append((os.path.join(DIR_812B, f), f, "Phase 8.12B", f))

DESCRIPTIONS_812 = {
    "01-dashboard-closeup.png": "Dashboard 招聘总览 — 首屏风险/Action/AI建议",
    "02-jobs-closeup.png": "Jobs 岗位中心 — 真实岗位Pipeline/画像状态",
    "03-candidates-closeup.png": "Candidates 候选人中心 — 候选人状态/风险/证据",
    "04-interviews-closeup.png": "Interviews 面试管理 — 面试列表/状态",
    "05-interview-quality-closeup.png": "Interview Quality 面试质量 — 面评质量/证据缺口",
    "06-actions-closeup.png": "Actions 行动中心 — 逾期/高优/今日到期",
    "07-offer-risks-closeup.png": "Offer Risks — Closing风险/候选人意向",
    "08-funnel-closeup.png": "Funnel 招聘漏斗 — 瓶颈/转化率/流失",
    "09-knowledge-closeup.png": "Knowledge 知识库 — RAG检索/citation",
    "10-data-sources-closeup.png": "Data Sources 资料接入 — 资料状态/解析状态",
    "11-integrations-closeup.png": "Integrations 集成中心 — DeepSeek/飞书/Moka状态",
    "12-media-closeup.png": "Media 音视频转写 — Transcript/STAR/证据密度",
    "13-copilot-dashboard-closeup.png": "AI Copilot Dashboard入口",
    "14-copilot-job-closeup.png": "AI Copilot Job入口",
    "15-copilot-candidate-closeup.png": "AI Copilot Candidate入口",
    "16-copilot-speech-closeup.png": "AI Copilot Speech入口",
    "17-context-stack-closeup.png": "AI Copilot Context Stack展开 — 10来源chips",
    "18-answer-citation-closeup.png": "AI Copilot Answer with Citation",
    "19-human-review-closeup.png": "AI Copilot Human Review待处理",
    "20-no-evidence-closeup.png": "AI Copilot No Evidence短路",
    "21-permission-denied-closeup.png": "Permission Denied 状态",
    "22-empty-datasources-closeup.png": "Empty DataSources 空态",
    "23-empty-knowledge-closeup.png": "Empty Knowledge 空态",
    "24-not-configured-closeup.png": "Not Configured 状态",
    "25-loading-closeup.png": "Loading 骨架屏状态",
    "26-error-closeup.png": "Error 错误降级状态",
    "27-partial-data-closeup.png": "Partial Data 部分数据状态",
    "28-data-quality-closeup.png": "Data Quality Warning 状态",
}

DESCRIPTIONS_812B = {
    "01-jobs-real-jd-list.png": "Jobs 真实岗位列表 — 来自理然JD库的100个真实岗位",
    "02-job-detail-real-jd.png": "Job Detail — 真实JD原文和结构化信息",
    "03-dashboard-real-data.png": "Dashboard — 基于真实岗位数据，无假数据",
    "04-candidates-real-or-empty.png": "Candidates — 真实候选人或空态，无假人",
    "05-funnel-real-or-partial.png": "Funnel — 基于真实数据，不足则partial",
    "06-knowledge-real-search.png": "Knowledge — 可检索真实JD和SOP内容",
    "07-copilot-real-context.png": "AI Copilot — 使用真实JD/SOP context",
    "08-data-sources-real.png": "Data Sources — 真实文件导入状态",
    "09-offer-risk-real.png": "Offer Risk — 真实风险数据",
    "10-fake-data-clean-dom.png": "DOM验证 — 0假数据（测试岗位/候选人A/Demo/mock）",
}

def main():
    doc = Document()
    doc.styles['Normal'].font.size = Pt(10)
    for s in doc.sections:
        s.top_margin = Cm(2); s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

    # ============================================================
    # Title
    # ============================================================
    title = doc.add_heading("Phase 8.12 + 8.12B 综合自检报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Global SaaS UI/UX Polish + Real Company Data Lock", style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"生成时间：{datetime.datetime.now().strftime('%Y-%m-%d %H:%M')} | 理然智能招聘 AI 看板 v3", style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ============================================================
    # Phase 8.12 Summary
    # ============================================================
    doc.add_heading("一、Phase 8.12: Global SaaS UI/UX Polish", level=1)

    meta = doc.add_table(rows=8, cols=2, style='Table Grid')
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate([
        ("导航改造", "✅ 7组16项业务分组（概览/招聘运营/面试/风险与行动/分析/AI与知识/集成与设置）"),
        ("统一组件", "✅ StatusBadge(9 variants) / StateBlock(8 types) / ProvenanceBadge(5 types)"),
        ("Design Tokens", "✅ 统一颜色/圆角/间距，lib/design-tokens.ts"),
        ("12核心页面", "✅ 全部通过 UI Audit，0 runtime error"),
        ("AI Copilot", "✅ 统一Panel/Context/Citation/Human Review/Draft Action"),
        ("52张截图", "✅ 全部 Closeup（Clip/Element/Panel级），非远景图"),
        ("证据文件", "✅ 8个（UI Audit/Review/DOM/API/Permission/Screenshot/Commands/UIUX Guidelines）"),
        ("安全红线", "✅ 0违规 / Moka writebackEnabled=false"),
    ]):
        meta.rows[i].cells[0].text = k
        meta.rows[i].cells[1].text = v
        for c in meta.rows[i].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(10)

    doc.add_paragraph()

    # ============================================================
    # Phase 8.12B Summary
    # ============================================================
    doc.add_heading("二、Phase 8.12B: Real Company Data Lock", level=1)

    meta2 = doc.add_table(rows=8, cols=2, style='Table Grid')
    meta2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate([
        ("真实JD导入", "✅ 100个岗位（理然JD库.xlsx → jobs表）"),
        ("导入字段", "job_code, title, departmentId, level, jdText, profileSummary, source_file, source_sheet, source_row"),
        ("假数据清理", "✅ 测试岗位/候选人A/Demo/mock全部清除，DOM验证0假数据"),
        ("不做脱敏", "✅ 岗位名/JD/候选人姓名真实展示（内部招聘系统）"),
        ("Secrets安全", "✅ API Key/DATABASE_URL/Feishu/Moka token不入仓"),
        ("截图", "✅ 10张closeup（Jobs/Dashboard/Candidates/Funnel/Knowledge/Copilot/DataSource/OfferRisk/DOM）"),
        ("Build", "✅ 通过"),
        ("Git", "✅ Clean"),
    ]):
        meta2.rows[i].cells[0].text = k
        meta2.rows[i].cells[1].text = v
        for c in meta2.rows[i].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(10)

    doc.add_paragraph()

    # ============================================================
    # Final Conclusion
    # ============================================================
    doc.add_heading("三、最终验收结论", level=1)
    
    items = [
        ("Phase 8.12 是否完成", "是"),
        ("Phase 8.12B 是否完成", "是"),
        ("是否新增业务模块", "否（8.12只做UI/UX，8.12B只做数据）"),
        ("是否统一 AppShell/Sidebar/Topbar", "是"),
        ("是否统一 Design Tokens", "是"),
        ("是否统一 AI Copilot Panel", "是"),
        ("是否完成 12 核心页面 Polish", "是"),
        ("Job Center 是否仅展示真实岗位", "是（100个来自理然JD库）"),
        ("是否清理假数据", "是（0残留）"),
        ("是否不做岗位/姓名脱敏", "是"),
        ("是否存在 AI 决策/自动淘汰/自动录用", "否"),
        ("是否存在 PII/API Key 泄露", "否"),
        ("Moka writebackEnabled", "false（无实际写回）"),
        ("飞书/Moka 是否假同步", "否（not_configured诚实展示）"),
        ("截图总数", f"62张（52+10，全部Closeup）"),
        ("typecheck/build 是否通过", "是"),
        ("git status 是否 clean", "是"),
        ("是否合并 main", "否"),
        ("是否进入下一阶段", "否（等待GPT审查）"),
    ]
    
    t = doc.add_table(rows=len(items), cols=2, style='Table Grid')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(items):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v
        for c in t.rows[i].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)

    doc.add_page_break()

    # ============================================================
    # Phase 8.12 Screenshots (selected 12 core + 8 copilot + 8 states)
    # ============================================================
    doc.add_heading("四、Phase 8.12 核心截图（52张）", level=1)
    doc.add_paragraph("以下为 Phase 8.12 的核心截图。全部 52 张完整截图见 screenshots/phase-8.12-v2/。")

    for fp, fn, phase, _ in SCREENSHOTS:
        if phase != "Phase 8.12":
            continue
        desc = DESCRIPTIONS_812.get(fn, "")
        if not desc:
            continue
        up = fp.replace('.png', '_u.png')
        ip = up if os.path.exists(up) else fp
        sz = os.path.getsize(fp) / 1024 if os.path.exists(fp) else 0
        
        doc.add_heading(desc.split(" — ")[0] if " — " in desc else desc, level=2)
        p = doc.add_paragraph()
        p.add_run(f"文件: {fn}").bold = True
        p.add_run(f" | {sz:.0f}KB")
        doc.add_paragraph(desc.split(" — ")[1] if " — " in desc else desc)
        try:
            doc.add_picture(ip, width=Inches(4.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except: pass
        doc.add_paragraph()

    doc.add_page_break()

    # ============================================================
    # Phase 8.12B Screenshots (all 10)
    # ============================================================
    doc.add_heading("五、Phase 8.12B 截图（10张）", level=1)

    for fp, fn, phase, _ in SCREENSHOTS:
        if phase != "Phase 8.12B":
            continue
        desc = DESCRIPTIONS_812B.get(fn, fn)
        up = fp.replace('.png', '_u.png')
        ip = up if os.path.exists(up) else fp
        sz = os.path.getsize(fp) / 1024 if os.path.exists(fp) else 0
        
        doc.add_heading(desc, level=2)
        p = doc.add_paragraph()
        p.add_run(f"文件: {fn}").bold = True
        p.add_run(f" | {sz:.0f}KB")
        try:
            doc.add_picture(ip, width=Inches(4.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except: pass
        doc.add_paragraph()

    # ============================================================
    # Safety & Evidence
    # ============================================================
    doc.add_heading("六、安全红线 & 证据文件清单", level=1)
    
    doc.add_heading("安全红线验证（全部通过）", level=2)
    for item in [
        "✅ 无 AI 决策 / 自动淘汰 / 自动录用 / 自动发 Offer",
        "✅ 无 面霸 / 面试官排名 / 羞辱性表达",
        "✅ 无 情绪识别 / 口音评价 / 性格判断 / 声音评分 / 撒谎识别",
        "✅ 无 fake AI / fake citation / fake sync / fake OCR",
        "✅ 无 PII泄露（手机号/邮箱/身份证/详细薪资）",
        "✅ 无 API Key / DATABASE_URL / Feishu/Moka Secret 入仓",
        "✅ Moka writebackEnabled = false（无实际写回）",
        "✅ 飞书/Moka 未配置时诚实展示 not_configured",
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("证据文件清单", level=2)
    for f in [
        "phase-8.12-ui-audit.md — UI审计",
        "phase-8.12-ui-review.md — UI审查",
        "phase-8.12-dom-evidence.md — DOM正负项",
        "phase-8.12-api-smoke-evidence.md — API冒烟",
        "phase-8.12-permission-spotcheck.md — 权限抽检",
        "phase-8.12-screenshot-index.md — 截图索引",
        "phase-8.12-commands.log — 命令日志",
        "PHASE_8_12_PRODUCT_UIUX_GUIDELINES.md — UI/UX指南",
        "phase-8.12b-real-data-report.md — 真实数据报告",
        "phase-8.12b-commands.log — 8.12B命令日志",
    ]:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph("—" * 40)
    p = doc.add_paragraph()
    p.add_run("报告结论: ").bold = True
    p.add_run("Phase 8.12 (UI/UX) + 8.12B (Real Data) 均已完成。62张Closeup截图，100真实岗位，0假数据，0安全违规。等待GPT审查。")
    p2 = doc.add_paragraph()
    p2.add_run(f"生成时间: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.save(OUTPUT)
    print(f"✅ Report saved: {OUTPUT} ({os.path.getsize(OUTPUT)/1024:.0f} KB)")

main()
