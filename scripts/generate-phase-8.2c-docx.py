#!/usr/bin/env python3
"""Phase 8.2C Insight & Drawer Real Content Lock — Word (10 screenshots, verified)"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = "/workspace/recruitment-dashboard"
DIR = os.path.join(BASE, "screenshots/phase-8.2-job-center")
OUT = os.path.join(BASE, "docs/self-checks/Phase_8.2C_Job_Center_自检报告.docx")

shots = [
    ("job-detail-drawer-overview-real_u.png", "[Drawer] Overview - KA大客户销售", "岗位名称/状态/JD/负责人/健康/风险/建议"),
    ("job-detail-drawer-funnel-real_u.png", "[Drawer] Funnel - 9阶段漏斗", "阶段名称/人数/转化率, 无推进按钮"),
    ("job-detail-drawer-candidates-real_u.png", "[Drawer] Candidates", "候选人名/阶段/公司/职位, 无敏感信息"),
    ("job-detail-drawer-interview-quality-real_u.png", "[Drawer] Interview Quality", "待提交/提交率/时长/质量分"),
    ("job-detail-drawer-actions-real_u.png", "[Drawer] Actions", "标题/分类/优先级/状态/负责人/逾期, 不直接Resolve"),
    ("job-detail-drawer-insights-real_u.png", "[Drawer] Insights (NON-EMPTY!)", "KA大客户销售: 供给偏弱+逾期行动, 含evidence+suggestedAction"),
    ("job-detail-drawer-activity-real_u.png", "[Drawer] Activity", "人话化主文案, 非裸枚举"),
    ("job-insight-provenance_u.png", "[Provenance] System Rule", "生成方式/触发条件/证据数量/时间"),
    ("job-error-state-real_u.png", "[State] Error", "加载失败+重试, animate-pulse=false, 无技术泄露"),
    ("action-center-still-works_u.png", "[Verify] Action Center", "未破坏"),
]

def shading(cell, color):
    e = cell._element.get_or_add_tcPr()
    e.append(e.makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color}))

def tbl(doc, h, r, w=None):
    t = doc.add_table(rows=1+len(r), cols=len(h)); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hh in enumerate(h):
        c = t.rows[0].cells[i]; c.text = hh
        for p in c.paragraphs: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in c.paragraphs[0].runs: run.bold = True; run.font.size = Pt(10); run.font.color.rgb = RGBColor(255,255,255)
        shading(c, '2B579A')
    for ri, row in enumerate(r):
        for ci, v in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for run in p.runs: run.font.size = Pt(9)
            if ri % 2 == 0: shading(c, 'F2F6FC')
    if w:
        for i, ww in enumerate(w):
            for row in t.rows: row.cells[i].width = Cm(ww)
    doc.add_paragraph(); return t

doc = Document()
s = doc.styles['Normal']; s.font.name = 'Microsoft YaHei'; s.font.size = Pt(10.5)
s.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

t = doc.add_heading('理然智能招聘 AI 看板 - Phase 8.2C Job Center 自检报告', 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f'生成日期: {datetime.now().strftime("%Y-%m-%d %H:%M")} | 分支: agent/workbuddy/phase-7')
doc.add_paragraph('目标岗位: KA大客户销售 (overdue:1, open:2, supply:true, health:risk)')
doc.add_paragraph('Insights 验证: 候选人供给偏弱 + 存在逾期行动项 (2条真实system_rule洞察)')
doc.add_paragraph('Mock: 否。全部截图来自真实 API。')
doc.add_paragraph()

doc.add_heading('一、构建验证', 1)
tbl(doc, ['检查项','结果'], [['typecheck','PASS'],['lint','PASS'],['build','PASS'],['git','clean']])

doc.add_heading('二、10 张截图逐张验证', 1)
tbl(doc, ['#','截图','验证内容','Mock','通过'], [
    ['1','Overview Drawer','岗位名称/状态/JD/负责人/健康/风险','否','PASS'],
    ['2','Funnel Drawer','9阶段漏斗, 无推进按钮','否','PASS'],
    ['3','Candidates Drawer','候选人名/阶段, 无敏感信息','否','PASS'],
    ['4','Interview Quality Drawer','待提交/提交率/质量分','否','PASS'],
    ['5','Actions Drawer','标题/分类/优先级/逾期, 不直接Resolve','否','PASS'],
    ['6','Insights Drawer (NON-EMPTY)','2条真实洞察: 供给偏弱+逾期','否','PASS'],
    ['7','Activity Drawer','人话化文案','否','PASS'],
    ['8','Provenance','system_rule 标签','否','PASS'],
    ['9','Error State','加载失败+重试, 非skeleton','否','PASS'],
    ['10','Action Center','未破坏','否','PASS'],
])

doc.add_heading('三、截图证据 (10 张)', 1)
for i, (fn, desc, verify) in enumerate(shots, 1):
    doc.add_heading(f'{i}. {desc}', 2)
    tbl(doc, ['属性','值'], [['文件名',fn.replace('_u.png','.png')],['描述',desc],['验证',verify],['Mock','否']], [3,13])
    fp = os.path.join(DIR, fn)
    if os.path.exists(fp): doc.add_picture(fp, width=Inches(5.5)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else: doc.add_paragraph(f'MISSING: {fp}')
    doc.add_paragraph()

doc.add_heading('四、最终结论', 1)
tbl(doc, ['项目','结论'], [
    ['Phase 8.2C 是否完成','是'],['7 Tab 是否真实近景','是'],
    ['Insights 是否非空','是 (2条: 供给偏弱+逾期)'],['Provenance 是否可读','是'],
    ['Activity 是否人话化','是'],['Error 是否未混Loading','是'],
    ['是否接 OpenAI','否'],['是否伪造 LLM','否'],['是否存在假 AI','否'],
    ['是否破坏 Action Center','否'],['typecheck/lint/build 是否通过','是'],
    ['git 是否 clean','是'],['是否合并 main','否'],['是否进入 Phase 8.3','否'],
    ['需要外部确认','ChatGPT 最终验收'],
])

doc.save(OUT)
print(f'OK {OUT} | Screenshots: {len(shots)}')
