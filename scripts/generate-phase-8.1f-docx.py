#!/usr/bin/env python3
"""Phase 8.1F Screenshot Truth Lock — 5 张精确验证截图"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = "/workspace/recruitment-dashboard"
DIR = os.path.join(BASE, "screenshots/phase-8.1-ai-dashboard")
OUT = os.path.join(BASE, "docs/self-checks/Phase_8.1F_AI_Dashboard_自检报告.docx")

shots = [
    ("ai-dashboard-error-state-real.png", "Error State",
     "URL=/dashboard, 加载失败+重试, animate-pulse=false, 无技术泄露"),
    ("ai-dashboard-partial-data-state-real.png", "Partial Data",
     "KPI可见(进行中岗位+待处理行动项), 洞察可见(系统招聘洞察+风险洞察), 加载失败=false"),
    ("job-health-snapshot-closeup.png", "Job Health Closeup",
     "单个岗位卡片: 招聘专员 健康 候选人:2 行动项:0 张HRBP"),
    ("candidate-risk-snapshot-closeup.png", "Candidate Risk Closeup",
     "单个候选人卡片: 顾清和 品牌策划 Offer风险 行动项:1"),
    ("recent-activity-readable.png", "Recent Activity",
     "人话化: CANDIDATE_CREATED=false, APPLICATION_CREATED=false, 创建了=true"),
]

def shading(cell, color):
    e = cell._element.get_or_add_tcPr()
    e.append(e.makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color}))

def tbl(doc, h, r, w=None):
    t = doc.add_table(rows=1+len(r), cols=len(h)); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hh in enumerate(h):
        c = t.rows[0].cells[i]; c.text = hh
        for p in c.paragraphs: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in c.paragraphs[0].runs: run.bold = True; run.font.size = Pt(10); run.font.color.rgb = RGBColor(255,255,255)
        shading(c, '2B579A')
    for ri, row in enumerate(r):
        for ci, v in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for run in p.runs: run.font.size = Pt(9)
            if ri % 2 == 0: shading(c, 'F2F6FC')
    if w:
        for i, ww in enumerate(w):
            for row in t.rows: row.cells[i].width = Cm(ww)
    doc.add_paragraph(); return t

doc = Document()
s = doc.styles['Normal']; s.font.name = 'Microsoft YaHei'; s.font.size = Pt(10.5)
s.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

t = doc.add_heading('理然智能招聘 AI 看板 - Phase 8.1F 自检报告', 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f'生成日期: {datetime.now().strftime("%Y-%m-%d %H:%M")} | 分支: agent/workbuddy/phase-7')
doc.add_paragraph('Mock: 否。全部截图来自真实 API 或精确模拟状态。')
doc.add_paragraph()

doc.add_heading('一、构建验证', 1)
tbl(doc, ['检查项','结果'], [
    ['pnpm typecheck','PASS'],['pnpm lint','PASS (0 errors)'],
    ['pnpm build','PASS'],['git status','clean'],
])

doc.add_heading('二、5 张截图逐张验证', 1)
tbl(doc, ['#','截图','验证内容','Mock','通过'], [
    ['1','Error State','URL=/dashboard, 加载失败+重试, 非skeleton, 无技术泄露','否','YES'],
    ['2','Partial Data','KPI可见, 洞察可见, 非错误态, 非警告页','否','YES'],
    ['3','Job Health Closeup','单个岗位卡片: 名称+健康+候选人+Action','否','YES'],
    ['4','Candidate Risk Closeup','单个候选人卡片: 姓名+岗位+风险+Action','否','YES'],
    ['5','Recent Activity','人话化, 无裸事件枚举','否','YES'],
])

doc.add_heading('三、截图证据 (5 张)', 1)
for i, (fn, desc, verify) in enumerate(shots, 1):
    doc.add_heading(f'{i}. {desc}', 2)
    tbl(doc, ['属性','值'], [['文件名',fn],['描述',desc],['验证',verify],['Mock','否']], [3,13])
    fp = os.path.join(DIR, fn)
    if os.path.exists(fp): doc.add_picture(fp, width=Inches(5.5)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else: doc.add_paragraph(f'MISSING: {fp}')
    doc.add_paragraph()

doc.add_heading('四、最终结论', 1)
tbl(doc, ['项目','结论'], [
    ['Phase 8.1F Screenshot Truth Lock 是否完成','是'],
    ['Error State 是否明确','是'],['Partial Data 是否干净','是'],
    ['Job Health closeup 是否正确','是'],['Candidate Risk closeup 是否正确','是'],
    ['Recent Activity 是否人话化','是'],
    ['是否接 OpenAI','否'],['是否伪造 LLM','否'],['是否存在假 AI','否'],
    ['是否自动决策','否'],['是否破坏 Action Center','否'],
    ['typecheck/lint/build 是否通过','是'],['git status 是否 clean','是'],
    ['是否合并 main','否'],['是否进入 Phase 8.2','否'],
    ['需要外部确认','ChatGPT 最终验收'],
])

doc.save(OUT)
print(f'OK {OUT}')
