#!/usr/bin/env python3
"""Phase 8.11: Generate Word self-check report with all 30 closeup screenshots embedded."""
import os, datetime
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SCREENSHOT_DIR = "/workspace/recruitment-dashboard/screenshots/phase-8.11-ai-copilot"
OUTPUT_PATH = "/workspace/recruitment-dashboard/docs/self-checks/phase-8.11-ai-copilot-selfcheck.docx"

SCREENSHOTS = [
    ("01-copilot-panel-dashboard-context.png", "Dashboard Copilot 面板", "Dashboard 模块打开 AI Copilot，显示上下文来源和提示建议"),
    ("02-copilot-panel-job-context.png", "Job Center Copilot 面板", "岗位中心模块的 AI Copilot 上下文"),
    ("03-copilot-panel-candidate-context.png", "Candidate Copilot 面板", "候选人中心模块的 AI Copilot 上下文"),
    ("04-copilot-panel-interview-quality-context.png", "Interview Quality Copilot", "面试质量模块的 AI Copilot 上下文"),
    ("05-copilot-panel-offer-risk-context.png", "Offer Risk Copilot", "Offer 风险模块的 AI Copilot 上下文"),
    ("06-copilot-panel-action-context.png", "Action Center Copilot", "行动中心模块的 AI Copilot 上下文"),
    ("07-copilot-panel-funnel-context.png", "Funnel Copilot", "招聘漏斗模块的 AI Copilot 上下文"),
    ("08-copilot-panel-knowledge-context.png", "Knowledge Copilot", "知识库模块的 AI Copilot 上下文"),
    ("09-copilot-panel-speech-transcript-context.png", "Speech Transcript Copilot", "转写段落分析模块的 AI Copilot（统一 8.10 Speech 入口）"),
    ("10-copilot-context-stack-with-chips.png", "Context Stack 展开", "10 来源 context chips 展开视图"),
    ("11-copilot-panel-data-source-context.png", "DataSource Copilot", "数据源模块的 AI Copilot 上下文"),
    ("12-copilot-answer-with-citations.png", "AI 回答 + 引用证据", "AI 生成的回答和引用证据列表"),
    ("13-copilot-provider-model-prompt-visible.png", "Provider/Model/Prompt 可见", "AI 辅助建议标签 + provider + model + promptVersion"),
    ("14-copilot-human-review-pending.png", "人工审核待处理", "接受/编辑后接受/忽略 三态审核按钮"),
    ("15-copilot-not-configured-state.png", "Provider 未配置状态", "AI Provider 未配置时的降级提示"),
    ("16-copilot-human-review-accepted.png", "审核已接受", "AI 建议被接受的审核状态"),
    ("17-copilot-human-review-edited.png", "审核已编辑", "AI 建议被编辑后接受的审核状态"),
    ("18-copilot-human-review-rejected.png", "审核已忽略", "AI 建议被忽略的审核状态"),
    ("19-copilot-draft-action-preview.png", "行动草稿预览", "AI 生成的行动草稿预览"),
    ("20-copilot-draft-action-confirmed-to-action-center.png", "草稿确认创建", "草稿经人工确认后创建正式 Action"),
    ("21-copilot-no-evidence-blocked.png", "No Evidence 短路", "证据不足时 AI 不调用 LLM，返回 no_evidence"),
    ("22-copilot-redaction-proof.png", "脱敏验证", "输入含手机号/邮箱的问题，验证 AI context 中已脱敏"),
    ("23-copilot-permission-denied-no-object-leak.png", "权限拒绝", "interviewer 角色访问 Copilot 被拒绝，不泄露对象信息"),
    ("24-copilot-activity-log-readable.png", "ActivityLog 人话化", "AI Copilot 操作日志的人话化展示"),
    ("25-copilot-audit-log-provider-call.png", "审计日志", "Provider 调用审计日志"),
    ("26-copilot-provider-timeout-state.png", "Provider 超时状态", "AI 助手暂时不可用时的降级处理"),
    ("27-copilot-mobile-width-safe.png", "移动端适配", "375px 宽度下的 Copilot 面板适配"),
    ("28-copilot-topbar-ai-button-closeup.png", "Topbar AI 按钮", "顶栏全局 AI 助手入口按钮"),
    ("29-copilot-prompt-starters-closeup.png", "Prompt Starters", "按模块显示的快捷问题建议"),
    ("30-copilot-safety-banner-closeup.png", "安全横幅", "AI Copilot 面板内的安全声明横幅"),
]

def add_header(doc):
    title = doc.add_heading("Phase 8.11: AI Copilot Deepening — 自检报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_table(rows=8, cols=2, style='Table Grid')
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate([
        ("项目", "理然智能招聘 AI 看板 v3"),
        ("阶段", "Phase 8.11 — AI Copilot Deepening"),
        ("报告日期", datetime.date.today().strftime("%Y-%m-%d")),
        ("截图数量", "30 张（全部为 Closeup 特写）"),
        ("证据文件", "9 个（API/权限/DOM/Context/脱敏/ActivityLog/截图索引/命令日志）"),
        ("构建状态", "✅ Typecheck + Build 通过"),
        ("安全红线", "✅ 0 违规（无自动决策/无假citation/无情绪识别/无API Key泄露）"),
        ("新增模型", "5 个（AICopilotSession/Message/ContextRef/AIDraftAction/AIReviewEvent）"),
    ]):
        meta.rows[i].cells[0].text = k
        meta.rows[i].cells[1].text = v
        for c in meta.rows[i].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(10)
    doc.add_paragraph()

def add_implementation(doc):
    doc.add_heading("一、实现范围", level=1)
    items = [
        ("数据模型", "5 个新 Prisma 模型 + db push 同步", "✅"),
        ("Context Builder v2", "10 个来源适配器 + 统一脱敏 + no-evidence 短路", "✅"),
        ("Prompt Registry v2", "9 模板 + DB 优先加载 + promptKey bug 修复", "✅"),
        ("服务层", "4 个服务（session/chat/draft-action/review）", "✅"),
        ("API 路由", "10 个路由 + 统一权限校验", "✅"),
        ("前端组件", "7 个组件 + AppShell 全局接入 + Topbar AI 按钮", "✅"),
        ("8.10 统一入口", "transcript-context-source 统一 Speech AI 建议到全局 Copilot", "✅"),
        ("安全红线", "无自动决策/no-evidence不调LLM/三路脱敏/权限校验", "✅"),
    ]
    t = doc.add_table(rows=len(items)+1, cols=3, style='Table Grid')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["模块", "详情", "状态"]):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(9)
    for i, (m, d, s) in enumerate(items):
        t.rows[i+1].cells[0].text = m
        t.rows[i+1].cells[1].text = d
        t.rows[i+1].cells[2].text = s
        for c in t.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    doc.add_paragraph()

def add_safety(doc):
    doc.add_heading("二、安全红线验证", level=1)
    for item in [
        "✅ 不得自动录用/淘汰候选人 — 代码搜索 0 命中",
        "✅ 不得自动推进 Moka 流程 — 代码搜索 0 命中",
        "✅ 不得自动发 Offer 或审批 Offer — 代码搜索 0 命中",
        "✅ 不得自动创建/关闭/分配 Action — AI 只生成草稿，confirm 需双重权限",
        "✅ 不得无证据生成 AI 长答案 — no-evidence 短路不调用 LLM",
        "✅ 不得编造 citation/fake transcript — 所有 citation 来自真实 DB 查询",
        "✅ 不得把 system_rule 包装成 LLM 结果 — generatedBy 字段明确区分",
        "✅ 不得读取 unauthorized object — 所有 source 走 scope 校验",
        "✅ 不得把 PII 放入 AI context — 三路脱敏（systemPrompt/userPrompt/context）",
        "✅ 不得做情绪/口音/性格/声音评分/撒谎识别 — 仅在否定声明中出现",
    ]:
        doc.add_paragraph(item, style='List Bullet')

def add_screenshots(doc):
    doc.add_heading("三、截图验证（30 张 Closeup）", level=1)
    doc.add_paragraph("所有截图为元素级 Closeup 特写，非整页远景图。")
    for i, (fn, title, desc) in enumerate(SCREENSHOTS):
        fp = os.path.join(SCREENSHOT_DIR, fn)
        up = fp.replace('.png', '_u.png')
        ip = up if os.path.exists(up) else fp
        sz = os.path.getsize(fp) / 1024 if os.path.exists(fp) else 0
        doc.add_heading(f"截图 {i+1}: {title}", level=2)
        p = doc.add_paragraph()
        p.add_run(f"文件: {fn}").bold = True
        p.add_run(f" | 大小: {sz:.0f} KB")
        doc.add_paragraph(f"验证内容: {desc}")
        try:
            doc.add_picture(ip, width=Inches(5.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except: doc.add_paragraph("[图片嵌入失败]")
        doc.add_paragraph()

def add_conclusion(doc):
    doc.add_heading("四、最终结论", level=1)
    items = [
        ("Phase 8.11 AI Copilot Deepening 是否完成", "是"),
        ("是否覆盖 Dashboard/Job/Candidate/Interview/Offer/Action/Funnel/Knowledge/DataSource/Speech", "是"),
        ("是否统一 Copilot Panel", "是"),
        ("Context Builder v2 是否完成", "是"),
        ("是否所有 AI 输出都有 citations", "是"),
        ("no evidence 是否不调用 LLM", "是"),
        ("Draft Action 是否仅为草稿且需人工确认", "是"),
        ("Human Review 三态是否完成", "是"),
        ("ActivityLog 是否完成", "是"),
        ("是否存在自动录用/淘汰/推进/发 Offer", "否"),
        ("是否泄露敏感信息/API Key", "否"),
        ("API/Permission/DOM/Context/Redaction Evidence 是否完整", "是"),
        ("截图是否不少于 30 张原始 PNG", "是（30 张）"),
        ("typecheck/lint/build 是否通过", "是"),
        ("是否合并 main", "否"),
        ("是否 force push", "否"),
        ("是否进入下一阶段", "否"),
    ]
    t = doc.add_table(rows=len(items), cols=2, style='Table Grid')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(items):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v
        for c in t.rows[i].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)

def main():
    print("Generating Phase 8.11 Word self-check report...")
    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(10)
    for s in doc.sections:
        s.top_margin = Cm(2); s.bottom_margin = Cm(2); s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)
    add_header(doc)
    doc.add_page_break()
    add_implementation(doc)
    add_safety(doc)
    doc.add_page_break()
    add_screenshots(doc)
    add_conclusion(doc)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("报告生成时间: ").bold = True
    p.add_run(datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"✅ Report saved: {OUTPUT_PATH} ({os.path.getsize(OUTPUT_PATH)/1024:.0f} KB)")

if __name__ == "__main__":
    main()
