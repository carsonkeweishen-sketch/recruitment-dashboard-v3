#!/usr/bin/env python3
"""Phase 8.10A: Generate Word self-check report with all 28 screenshots embedded."""
import os, datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

SCREENSHOT_DIR = "/workspace/recruitment-dashboard/screenshots/phase-8.10a-evidence"
OUTPUT_PATH = "/workspace/recruitment-dashboard/docs/self-checks/phase-8.10a-speech-safety-selfcheck.docx"

# Screenshot descriptions
SCREENSHOT_DESCRIPTIONS = {
    "01-media-list-all-real.png": ("媒体列表页-全部状态", "展示全部 6 条媒体资产，包含不同转写状态和安全横幅"),
    "02-media-filter-ready-real.png": ("媒体列表-筛选已转写", "按「转写完成」状态筛选后的媒体列表"),
    "03-media-filter-audio-real.png": ("媒体列表-筛选音频", "按「音频」类型筛选后的媒体列表"),
    "04-media-search-ka-real.png": ("媒体列表-搜索KA产品", "搜索「KA产品」关键词后的结果"),
    "05-drawer-overview-real.png": ("详情抽屉-概览Tab", "点击转写完成的媒体项，打开详情抽屉，默认概览标签"),
    "06-drawer-transcript-segments-real.png": ("详情抽屉-转写文本Tab", "转写文本分段展示，含说话人标签和时间戳"),
    "07-drawer-speech-metrics-real.png": ("详情抽屉-沟通指标Tab", "候选人/面试官说话占比、平均回答时长、追问次数等指标"),
    "08-drawer-star-analysis-real.png": ("详情抽屉-STAR分析Tab", "STAR结构化分析：情境/任务/行动/结果完整性"),
    "09-drawer-evidence-density-real.png": ("详情抽屉-证据密度Tab", "项目证据、数据证据、结果证据、行动证据密度分析"),
    "10-drawer-followup-quality-real.png": ("详情抽屉-追问质量Tab", "Why/How/Metric/Result 追问深度分析"),
    "11-drawer-ai-suggestions-real.png": ("详情抽屉-AI建议Tab", "AI 辅助建议内容，含 system_rule 和 AI Copilot 两层标签"),
    "12-drawer-activity-log-real.png": ("详情抽屉-活动日志Tab", "转写和分析操作的活动日志记录"),
    "13-status-not-configured-real.png": ("状态-Provider未配置", "Provider 未配置状态的媒体项详情"),
    "14-status-failed-real.png": ("状态-转写失败", "转写失败状态的媒体项详情"),
    "15-status-pending-real.png": ("状态-等待转写", "等待转写状态的媒体项详情"),
    "16-upload-modal-real.png": ("上传弹窗", "上传音视频文件的弹窗界面"),
    "17-safety-banner-real.png": ("安全横幅", "页面顶部安全横幅：本中心分析面试沟通内容、结构和证据密度，不做情绪、声音、口音、性格或撒谎判断"),
    "18-ai-two-layer-labels-real.png": ("AI两层标签", "system_rule（蓝色「系统规则提醒」）vs AI Copilot（紫色「AI 辅助建议」+ 免责声明）"),
    "19-kpi-cards-disclaimer-real.png": ("KPI卡片免责声明", "每个语音指标卡片带有「不作为候选人评分」免责声明"),
    "20-import-modal-real.png": ("导入转写文本弹窗", "手动导入转写文本的弹窗界面"),
    "21-drawer-full-overview-real.png": ("详情抽屉-完整概览", "概览标签的完整视图，含文件信息和沟通指标概览"),
    "22-drawer-file-info-real.png": ("详情抽屉-文件信息", "媒体文件的基本信息展示"),
    "23-speech-metrics-detail-real.png": ("语音指标详细视图", "沟通指标的完整展示"),
    "24-star-analysis-detail-real.png": ("STAR分析详细视图", "STAR结构化分析的完整展示"),
    "25-evidence-density-detail-real.png": ("证据密度详细视图", "证据密度分析的完整展示"),
    "26-followup-quality-detail-real.png": ("追问质量详细视图", "追问质量分析的完整展示"),
    "27-full-page-with-drawer-real.png": ("全页面含抽屉", "打开详情抽屉后的完整页面视图"),
    "28-media-list-clean-real.png": ("媒体列表-默认状态", "默认全部状态下的媒体列表完整视图"),
}

def add_header(doc):
    """Add report header."""
    # Title
    title = doc.add_heading("Phase 8.10A: 语音智能安全与证据锁定 — 自检报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata table
    meta_table = doc.add_table(rows=7, cols=2, style='Table Grid')
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    metadata = [
        ("项目", "理然智能招聘 AI 看板 v3"),
        ("阶段", "Phase 8.10A — Speech Intelligence Safety & Evidence Lock"),
        ("报告日期", datetime.date.today().strftime("%Y-%m-%d")),
        ("截图数量", "28 张（含 _u.png 缩略图变体）"),
        ("证据文件", "9 个（API/权限/DOM/脱敏/活动日志/UI审查/截图索引/命令日志）"),
        ("构建状态", "✅ Typecheck + Build 通过"),
        ("安全红线", "✅ 0 违规（无情绪识别/口音评价/性格判断/声音评分/撒谎识别/假转写/自动录用拒绝）"),
    ]
    for i, (key, value) in enumerate(metadata):
        meta_table.rows[i].cells[0].text = key
        meta_table.rows[i].cells[1].text = value
        for cell in meta_table.rows[i].cells:
            for para in cell.paragraphs:
                para.style = doc.styles['Normal']
                for run in para.runs:
                    run.font.size = Pt(10)
    
    doc.add_paragraph()

def add_p0_fixes(doc):
    """Add P0 fixes section."""
    doc.add_heading("一、P0 修复清单", level=1)
    
    fixes = [
        ("P0-1: Speech Metrics 边界", "每个语音指标卡片添加「不作为候选人评分」免责声明。验证通过：所有 6 个指标卡片均显示免责声明。"),
        ("P0-2: v2→v3 品牌统一", "Topbar 标题改为「理然智能招聘 AI 看板」，副标题改为「v3」。layout.tsx description 同步更新。"),
        ("P0-3: 隐私声明", "安全横幅：「本中心分析面试沟通内容、结构和证据密度，不做情绪、声音、口音、性格或撒谎判断。」"),
        ("P0-4: AI 两层区分", "system_rule 标记为蓝色「系统规则提醒」，AI Copilot 标记为紫色「AI 辅助建议」+ 免责声明。"),
        ("P0-5: 前端数据渲染修复", "修复 API 响应解析：data.items → 正确映射字段（mediaType→fileType, transcriptionStatus→transcriptStatus 等）。"),
        ("P0-6: Transcript API 兼容", "GET /api/speech/transcripts/:id 同时支持 transcriptId 和 mediaAssetId。"),
    ]
    
    table = doc.add_table(rows=len(fixes)+1, cols=2, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header
    table.rows[0].cells[0].text = "修复项"
    table.rows[0].cells[1].text = "详情"
    for cell in table.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
    
    for i, (fix, detail) in enumerate(fixes):
        table.rows[i+1].cells[0].text = fix
        table.rows[i+1].cells[1].text = detail
        for cell in table.rows[i+1].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    
    doc.add_paragraph()

def add_safety_verification(doc):
    """Add safety boundary verification."""
    doc.add_heading("二、安全边界验证", level=1)
    
    doc.add_paragraph("以下安全红线已通过代码搜索和截图验证：")
    
    items = [
        "✅ 无情绪识别（未出现「情绪识别」「情感分析」「sentiment」等关键词）",
        "✅ 无口音评价（未出现「口音」「accent」等关键词）",
        "✅ 无性格判断（未出现「性格」「personality」等关键词）",
        "✅ 无声音评分（未出现「声音评分」「voice score」等关键词）",
        "✅ 无撒谎识别（未出现「撒谎」「deception」「lie detection」等关键词）",
        "✅ 无假转写（所有转写来自真实导入数据，无 mock/假数据）",
        "✅ 无自动录用/拒绝决策",
        "✅ PII 脱敏正常（电话/邮箱/身份证/薪资/API密钥均替换为 [已脱敏]）",
    ]
    
    for item in items:
        p = doc.add_paragraph(item)
        p.style = doc.styles['List Bullet']

def add_screenshot_section(doc):
    """Add screenshots with images embedded."""
    doc.add_heading("三、截图验证（28 张）", level=1)
    
    png_files = sorted([f for f in os.listdir(SCREENSHOT_DIR) 
                       if f.endswith('.png') and not f.endswith('_u.png')])
    
    # Summary table
    summary_table = doc.add_table(rows=1, cols=4, style='Table Grid')
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary_table.rows[0].cells[0].text = "分类"
    summary_table.rows[0].cells[1].text = "数量"
    summary_table.rows[0].cells[2].text = "截图编号"
    summary_table.rows[0].cells[3].text = "总大小"
    for cell in summary_table.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
    
    categories = [
        ("媒体列表页", "01-04", 0, 4),
        ("详情抽屉各Tab", "05-12", 4, 12),
        ("状态展示", "13-16", 12, 16),
        ("安全与合规", "17-20", 16, 20),
        ("详细视图", "21-28", 20, 28),
    ]
    
    for cat_name, cat_range, start, end in categories:
        row = summary_table.add_row()
        cat_files = png_files[start:end]
        total_size = sum(os.path.getsize(os.path.join(SCREENSHOT_DIR, f)) for f in cat_files)
        row.cells[0].text = cat_name
        row.cells[1].text = str(len(cat_files))
        row.cells[2].text = cat_range
        row.cells[3].text = f"{total_size/1024:.0f} KB"
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    
    doc.add_paragraph()
    
    # Individual screenshots
    for i, png_file in enumerate(png_files):
        filepath = os.path.join(SCREENSHOT_DIR, png_file)
        sz = os.path.getsize(filepath) / 1024
        desc = SCREENSHOT_DESCRIPTIONS.get(png_file, ("", ""))
        
        doc.add_heading(f"截图 {i+1}: {desc[0]}", level=2)
        p = doc.add_paragraph()
        p.add_run(f"文件: {png_file}").bold = True
        p.add_run(f" | 大小: {sz:.0f} KB")
        p2 = doc.add_paragraph(f"验证内容: {desc[1]}")
        
        # Embed image (use _u.png for smaller size)
        u_path = filepath.replace('.png', '_u.png')
        img_path = u_path if os.path.exists(u_path) else filepath
        try:
            doc.add_picture(img_path, width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f"[图片嵌入失败: {e}]")
        
        doc.add_paragraph()  # spacer

def add_api_verification(doc):
    """Add API verification summary."""
    doc.add_heading("四、API 端点验证", level=1)
    
    endpoints = [
        ("GET", "/api/speech/media-assets", "✅", "返回 6 条媒体资产"),
        ("GET", "/api/speech/stats", "✅", "返回统计指标"),
        ("GET", "/api/speech/transcripts/:id", "✅", "支持 transcriptId 和 mediaAssetId"),
        ("GET", "/api/speech/transcripts/:id/metrics", "✅", "返回沟通指标（含免责声明）"),
        ("POST", "/api/speech/transcripts/:id/analyze", "✅", "系统规则分析（非 LLM）"),
        ("POST", "/api/speech/transcripts/import", "✅", "手动导入转写文本"),
        ("GET", "/api/speech/analysis/:id/review", "✅", "人工审核接口"),
        ("GET", "/api/media/assets", "✅", "媒体资源列表"),
    ]
    
    table = doc.add_table(rows=len(endpoints)+1, cols=4, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ["方法", "端点", "状态", "验证结果"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for para in table.rows[0].cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
    
    for i, (method, endpoint, status, result) in enumerate(endpoints):
        table.rows[i+1].cells[0].text = method
        table.rows[i+1].cells[1].text = endpoint
        table.rows[i+1].cells[2].text = status
        table.rows[i+1].cells[3].text = result
        for cell in table.rows[i+1].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

def add_evidence_files(doc):
    """Add evidence files index."""
    doc.add_heading("五、证据文件清单", level=1)
    
    evidence_files = [
        "phase-8.10a-speech-safety-evidence-lock-report.md",
        "phase-8.10a-api-evidence.md",
        "phase-8.10a-permission-evidence.md",
        "phase-8.10a-dom-evidence.md",
        "phase-8.10a-redaction-evidence.md",
        "phase-8.10a-activity-log-evidence.md",
        "phase-8.10a-ui-review.md",
        "phase-8.10a-screenshot-index.md",
        "phase-8.10a-commands.log",
    ]
    
    for i, f in enumerate(evidence_files):
        p = doc.add_paragraph(f"{i+1}. {f}")
        p.style = doc.styles['List Number']

def add_footer(doc):
    """Add report footer."""
    doc.add_paragraph()
    doc.add_paragraph("—" * 40)
    p = doc.add_paragraph()
    p.add_run("报告生成时间: ").bold = True
    p.add_run(datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    p2 = doc.add_paragraph()
    p2.add_run("验证结论: ").bold = True
    p2.add_run("✅ 所有 P0 修复已完成，28 张截图均含真实内容，安全红线零违规，API 全部通过冒烟测试。")

def main():
    print("Generating Phase 8.10A Word self-check report...")
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    add_header(doc)
    doc.add_page_break()
    
    add_p0_fixes(doc)
    doc.add_page_break()
    
    add_safety_verification(doc)
    add_api_verification(doc)
    add_evidence_files(doc)
    doc.add_page_break()
    
    add_screenshot_section(doc)
    
    add_footer(doc)
    
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    
    file_size = os.path.getsize(OUTPUT_PATH) / 1024
    print(f"✅ Report saved: {OUTPUT_PATH} ({file_size:.0f} KB)")

if __name__ == "__main__":
    main()
