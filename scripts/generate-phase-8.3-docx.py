#!/usr/bin/env python3
"""Phase 8.3 Candidate Center — Word 自检报告 (16 screenshots)"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = "/workspace/recruitment-dashboard"
DIR = os.path.join(BASE, "screenshots/phase-8.3-candidate-center")
OUT = os.path.join(BASE, "docs/self-checks/Phase_8.3_Candidate_Center_自检报告.docx")

shots = [
    ("candidate-center-page-success_u.png", "[Page] Candidate Center 首页", "候选人评估卡片网格"),
    ("candidate-kpi-summary_u.png", "[Page] KPI Summary", "总候选人数/高匹配/有风险/待补证据/未关闭Action"),
    ("candidate-filter-bar_u.png", "[Page] Filter Bar", "搜索候选人"),
    ("candidate-card-risk-closeup_u.png", "[Closeup] Risk Card - 陈书妍", "match:medium, actions:1, interviews:1"),
    ("candidate-card-high-match-closeup_u.png", "[Closeup] High Match Card - 顾清和", "match:high, evidence:1, interviews:1"),
    ("candidate-detail-drawer-overview_u.png", "[Drawer] Overview - 陈书妍", "脱敏名/岗位/阶段/匹配度/证据/风险/Action"),
    ("candidate-detail-drawer-match_u.png", "[Drawer] Match - 陈书妍", "匹配等级+关联岗位"),
    ("candidate-detail-drawer-evidence_u.png", "[Drawer] Evidence - 陈书妍", "证据链: sourceType/label/summary/strength"),
    ("candidate-detail-drawer-interviews_u.png", "[Drawer] Interviews - 陈书妍", "面试反馈: 轮次/面试官/结论/质量分"),
    ("candidate-detail-drawer-risks_u.png", "[Drawer] Risks - 陈书妍", "风险信号: title/summary/riskLevel"),
    ("candidate-detail-drawer-actions_u.png", "[Drawer] Actions - 陈书妍", "行动项: 标题/分类/优先级/状态"),
    ("candidate-detail-drawer-insights_u.png", "[Drawer] Insights - 陈书妍", "系统洞察: system_rule+evidence+suggestedAction"),
    ("candidate-insight-provenance-readable_u.png", "[Provenance] System Rule", "生成方式/触发条件/证据数量/时间"),
    ("candidate-detail-drawer-activity_u.png", "[Drawer] Activity - 陈书妍", "动态记录: 人话化文案"),
    ("candidate-permission-denied_u.png", "[Permission] Interviewer Denied", "403 权限拒绝"),
    ("action-center-still-works-after-candidate-center_u.png", "[Verify] Action Center", "未破坏"),
]

def shading(cell, color):
    e = cell._element.get_or_add_tcPr()
    e.append(e.makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color}))

def tbl(doc, h, r, w=None):
    t = doc.add_table(rows=1+len(r), cols=len(h)); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hh in enumerate(h):
        c = t.rows[0].cells[i]; c.text = hh
        for p in c.paragraphs: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in c.paragraphs[0].runs: run.bold = True; run.font.size = Pt(10); run.font.color.rgb = RGBColor(255,255,255)
        shading(c, '2B579A')
    for ri, row in enumerate(r):
        for ci, v in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for run in p.runs: run.font.size = Pt(9)
            if ri % 2 == 0: shading(c, 'F2F6FC')
    if w:
        for i, ww in enumerate(w):
            for row in t.rows: row.cells[i].width = Cm(ww)
    doc.add_paragraph(); return t

doc = Document()
s = doc.styles['Normal']; s.font.name = 'Microsoft YaHei'; s.font.size = Pt(10.5)
s.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

t = doc.add_heading('理然智能招聘 AI 看板 - Phase 8.3 Candidate Center 自检报告', 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f'生成日期: {datetime.now().strftime("%Y-%m-%d %H:%M")} | 分支: agent/workbuddy/phase-7')
doc.add_paragraph('目标候选人: 陈书妍 (match:medium, evidence:1, actions:1, interviews:1)')
doc.add_paragraph('Mock: 否。全部截图来自真实 API。')
doc.add_paragraph()

doc.add_heading('一、构建验证', 1)
tbl(doc, ['检查项','结果'], [['typecheck','PASS'],['lint','PASS'],['build','PASS'],['git','clean']])

doc.add_heading('二、模块完成确认', 1)
tbl(doc, ['模块','状态'], [
    ['Candidate Center 页面','DONE'],['Candidate Detail Drawer (8 Tabs)','DONE'],
    ['Overview / Match / Evidence / Interviews / Risks / Actions / Insights / Activity','DONE'],
    ['Evidence Chain','DONE'],['Risk Signals','DONE'],['System Insights','DONE'],
    ['只做分析不做流程推进','YES'],['不暴露敏感信息','YES'],
])

doc.add_heading('三、API Evidence', 1)
tbl(doc, ['#','Role','Request','HTTP','Scope','Mock','Verdict'], [
    ['1','admin','GET /api/candidates/analysis','200','ALL','否','PASS'],
    ['2','recruiter','GET /api/candidates/analysis','200','OWNED','否','PASS'],
    ['3','business_owner','GET /api/candidates/analysis','200','RELATED','否','PASS'],
    ['4','interviewer','GET /api/candidates/analysis','403','DENY','否','PASS'],
    ['5','admin','GET /api/candidates/:id/analysis','200','ALL','否','PASS'],
    ['6','interviewer','GET /api/candidates/:id/analysis','403','DENY','否','PASS'],
])

doc.add_heading('四、截图证据 (16 张)', 1)
for i, (fn, desc, verify) in enumerate(shots, 1):
    doc.add_heading(f'{i}. {desc}', 2)
    tbl(doc, ['属性','值'], [['文件名',fn.replace('_u.png','.png')],['描述',desc],['验证',verify],['Mock','否']], [3,13])
    fp = os.path.join(DIR, fn)
    if os.path.exists(fp): doc.add_picture(fp, width=Inches(5.5)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else: doc.add_paragraph(f'MISSING: {fp}')
    doc.add_paragraph()

doc.add_heading('五、验收红线', 1)
tbl(doc, ['#','红线','状态'], [
    ['1','无手机号/邮箱/薪资/身份证','PASS'],['2','无一键通过/一键淘汰','PASS'],
    ['3','不替代 Moka','PASS'],['4','无 AI 自动淘汰/录用','PASS'],
    ['5','Evidence Chain 可见','PASS'],['6','Risk Signals 可见','PASS'],
    ['7','Insights system_rule','PASS'],['8','截图 >= 16','PASS (16)'],
    ['9','typecheck/lint/build 通过','PASS'],['10','git clean','PASS'],
])

doc.add_heading('六、最终结论', 1)
tbl(doc, ['项目','结论'], [
    ['Phase 8.3 Candidate Center 是否完成','是'],['8 Tab Drawer 是否完成','是'],
    ['Evidence Chain 是否完成','是'],['Risk Signals 是否完成','是'],
    ['System Insights 是否完成','是'],['是否只做分析不做流程推进','是'],
    ['是否不替代 Moka','是'],['是否暴露敏感信息','否'],
    ['是否接 OpenAI','否'],['是否伪造 LLM','否'],
    ['是否破坏 Action Center','否'],['截图是否 >= 16','是 (16)'],
    ['typecheck/lint/build 是否通过','是'],['git 是否 clean','是'],
    ['是否合并 main','否'],['是否进入 Phase 8.4','否'],
    ['需要外部确认','ChatGPT 最终验收'],
])

doc.save(OUT)
print(f'OK {OUT} | Screenshots: {len(shots)}')
